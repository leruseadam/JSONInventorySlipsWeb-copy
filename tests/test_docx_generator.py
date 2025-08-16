"""
Tests for the DocxGenerator class
"""

import os
import unittest
from docx import Document
from datetime import datetime
from docx.enum.section import WD_ORIENT
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from src.utils.docgen import DocxGenerator

class TestDocxGenerator(unittest.TestCase):
    def setUp(self):
        """Set up test environment"""
        self.generator = DocxGenerator()
        self.test_output_dir = os.path.join(os.path.dirname(__file__), 'test_output')
        os.makedirs(self.test_output_dir, exist_ok=True)
        
    def tearDown(self):
        """Clean up test files"""
        for filename in os.listdir(self.test_output_dir):
            filepath = os.path.join(self.test_output_dir, filename)
            try:
                os.remove(filepath)
            except:
                pass
                
    def test_create_document(self):
        """Test creating a new document"""
        self.generator.create_document(landscape=True)
        
        # Verify document was created
        self.assertIsNotNone(self.generator.doc)
        
        # Verify landscape orientation
        section = self.generator.doc.sections[0]
        self.assertEqual(section.orientation, WD_ORIENT.LANDSCAPE)
        self.assertEqual(float(section.page_width), float(Inches(11)))
        self.assertEqual(float(section.page_height), float(Inches(8.5)))
        
    def test_add_heading(self):
        """Test adding a heading"""
        self.generator.create_document()
        self.generator.add_heading("Test Heading", size=16, bold=True)
        
        # Verify heading properties
        para = self.generator.doc.paragraphs[0]
        run = para.runs[0]
        self.assertEqual(run.text, "Test Heading")
        self.assertTrue(run.bold)
        self.assertEqual(run.font.size.pt, 16)
        self.assertEqual(para.alignment, WD_ALIGN_PARAGRAPH.CENTER)
        
    def test_generate_inventory_slip(self):
        """Test generating a complete inventory slip"""
        records = [
            {
                'Product Name*': 'Test Product 1',
                'Barcode*': '12345',
                'Quantity Received*': '10',
                'Vendor': 'Test Vendor',
                'Accepted Date': '2023-01-01'
            },
            {
                'Product Name*': 'Test Product 2',
                'Barcode*': '67890',
                'Quantity Received*': '20',
                'Vendor': 'Test Vendor',
                'Accepted Date': '2023-01-01'
            }
        ]
        
        # Generate test document
        self.generator.generate_inventory_slip(
            records=records,
            vendor_name='Test Vendor',
            date='2023-01-01',
            rows_per_page=20
        )
        
        # Save test document
        test_file = os.path.join(self.test_output_dir, 'test_slip.docx')
        self.generator.save(test_file)
        
        # Verify file was created
        self.assertTrue(os.path.exists(test_file))
        
        # Load and verify document content
        doc = Document(test_file)
        
        # Verify title
        self.assertEqual(doc.paragraphs[0].text, "Order Sheet")
        
        # Verify info line
        self.assertEqual(doc.paragraphs[1].text, "Date: 2023-01-01    Vendor: Test Vendor")
        
        # Verify table exists
        self.assertGreater(len(doc.tables), 0)
        table = doc.tables[0]
        
        # Verify headers
        expected_headers = ['Product Name', 'Barcode', 'Quantity', 'Vendor', 'Accepted Date']
        for idx, header in enumerate(expected_headers):
            cell_text = table.cell(0, idx).text.strip()
            self.assertEqual(cell_text, header)
            
        # Verify data rows
        self.assertEqual(table.cell(1, 0).text.strip(), "Test Product 1")
        self.assertEqual(table.cell(1, 1).text.strip(), "12345")
        self.assertEqual(table.cell(2, 0).text.strip(), "Test Product 2")
        self.assertEqual(table.cell(2, 1).text.strip(), "67890")
        
if __name__ == '__main__':
    unittest.main()
