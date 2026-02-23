import os
import sys
import json
from docx import Document
from docx.shared import Pt
from docxcompose.composer import Composer
from io import BytesIO
from docxtpl import DocxTemplate
import datetime
from collections import deque

def chunk_records(records, chunk_size=4):
    """Split records into chunks of specified size"""
    for i in range(0, len(records), chunk_size):
        yield records[i:i + chunk_size]

def adjust_table_font_sizes(doc_path):
    """
    Post-process a DOCX file to dynamically adjust font size inside table cells based on thresholds.
    """
    thresholds = [
        (30, 12),   # <=30 chars → 12pt
        (45, 10),   # <=45 chars → 10pt
        (60, 8),    # <=60 chars → 8pt
        (float('inf'), 7)  # >60 chars → 7pt
    ]

    def get_font_size(text_len):
        for limit, size in thresholds:
            if text_len <= limit:
                return size
        return 7  # Fallback

    doc = Document(doc_path)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text.strip()
                    if not text:
                        continue

                    # If line is Product Name (first line), force 10pt
                    if paragraph == cell.paragraphs[0]:
                        font_size = 10
                    else:
                        font_size = get_font_size(len(text))

                    for run in paragraph.runs:
                        run.font.size = Pt(font_size)

    doc.save(doc_path)

def open_file(path):
    """Open a file using the system's default application"""
    try:
        if sys.platform == "darwin":
            os.system(f'open "{path}"')
        elif sys.platform == "win32":
            os.startfile(path)
        else:
            os.system(f'xdg-open "{path}"')
    except Exception as e:
        print(f"Error opening file: {e}")

def format_json_text(text):
    """Format JSON text for better readability"""
    try:
        if not text.strip():
            return text
        
        parsed = json.loads(text)
        return json.dumps(parsed, indent=2)
    except json.JSONDecodeError:
        return text
    except Exception:
        return text

def run_full_process_inventory_slips(selected_df, config, status_callback=None, progress_callback=None):
    """Process and generate inventory slips"""
    if selected_df.empty:
        if status_callback:
            status_callback("Error: No data selected.")
        return False, "No data selected."

    try:
        # Get settings from config
        items_per_page = int(config['SETTINGS'].get('items_per_page', '4'))
        template_path = config['PATHS'].get('template_path')
        output_dir = config['PATHS'].get('output_dir')
        
        # Ensure output directory exists
        if not os.path.exists(output_dir):
            try:
                os.makedirs(output_dir)
            except Exception as e:
                return False, f"Failed to create output directory: {e}"
        
        if status_callback:
            status_callback("Processing data...")
        
        records = selected_df.to_dict(orient="records")
        pages = []

        # Create queues per logical slot so each slot holds one product type
        def map_type_to_slot(ptype):
            if not ptype:
                return None
            p = str(ptype).lower()
            if 'edibl' in p or 'edible' in p:
                return 1
            if 'bever' in p or 'drink' in p:
                return 2
            if 'flower' in p or 'bud' in p or 'floral' in p:
                return 3
            if 'concentrate' in p or 'wax' in p or 'shatter' in p or 'oil' in p:
                return 4
            return None

        # Build queues for slots 1..items_per_page
        slot_queues = {i: deque() for i in range(1, items_per_page + 1)}
        misc_queue = deque()

        for rec in records:
            ptype = rec.get('Product Type*', rec.get('Inventory Type', ''))
            slot = map_type_to_slot(ptype)
            if slot and slot in slot_queues:
                slot_queues[slot].append(rec)
            else:
                misc_queue.append(rec)

        # Append misc records into any empty slot queues in round-robin fashion
        # This ensures records that don't match known types still get placed.
        misc_idx = 1
        while misc_queue:
            placed = False
            for i in range(1, items_per_page + 1):
                if misc_idx > len(misc_queue):
                    misc_idx = 1
                if len(slot_queues[i]) == 0:
                    slot_queues[i].append(misc_queue.popleft())
                    placed = True
                    break
            if not placed:
                # All slots have at least one item; break so we'll distribute in paging step
                break

        # Determine how many pages required by the largest queue
        max_queue_len = max((len(q) for q in slot_queues.values()), default=0)
        # If misc_queue still has items, they'll create additional pages
        if misc_queue:
            # total pages is max of existing queues plus ceil(misc/slots)
            import math
            extra_pages = math.ceil(len(misc_queue) / items_per_page)
            total_pages = max_queue_len + extra_pages
        else:
            total_pages = max_queue_len

        if total_pages == 0 and records:
            total_pages = 1

        # Progress calculation
        current_chunk = 0

        # Build pages by pulling one item per slot per page (if available)
        for page_idx in range(total_pages):
            current_chunk += 1
            if progress_callback:
                progress_value = (current_chunk / max(total_pages,1)) * 50  # First half of progress
                progress_callback(int(progress_value))

            if status_callback:
                status_callback(f"Generating page {current_chunk} of {total_pages}...")

            try:
                tpl = DocxTemplate(template_path)
                context = {}

                # For each slot index, pop next record if available
                for slot_num in range(1, items_per_page + 1):
                    rec = None
                    if slot_queues.get(slot_num) and len(slot_queues[slot_num]) > 0:
                        rec = slot_queues[slot_num].popleft()
                    elif misc_queue:
                        rec = misc_queue.popleft()

                    if rec:
                        product_name = rec.get("Product Name*", "")
                        barcode = rec.get("Barcode*", "")
                        qty = rec.get("Quantity Received*", rec.get("Quantity*", ""))
                        try:
                            qty = int(float(qty))
                        except (ValueError, TypeError):
                            qty = ""

                        context[f"Label{slot_num}"] = {
                            "ProductName": product_name,
                            "Barcode": barcode,
                            "AcceptedDate": rec.get("Accepted Date", ""),
                            "QuantityReceived": qty,
                            "Vendor": rec.get("Vendor", ""),
                            "StrainName": rec.get("Strain Name", ""),
                            "ProductType": rec.get("Product Type*", rec.get("Inventory Type", "")),
                            "THCContent": rec.get("THC Content", ""),
                            "CBDContent": rec.get("CBD Content", "")
                        }
                    else:
                        context[f"Label{slot_num}"] = {
                            "ProductName": "",
                            "Barcode": "",
                            "AcceptedDate": "",
                            "QuantityReceived": "",
                            "Vendor": "",
                            "StrainName": "",
                            "ProductType": "",
                            "THCContent": "",
                            "CBDContent": ""
                        }

                tpl.render(context)
                buf = BytesIO()
                tpl.save(buf)
                pages.append(Document(buf))

            except Exception as e:
                return False, f"Error generating page {current_chunk}: {e}"
        
        if not pages:
            return False, "No documents generated."
        
        if status_callback:
            status_callback("Combining pages...")
        
        master = pages[0]
        composer = Composer(master)
        for i, doc in enumerate(pages[1:]):
            if progress_callback:
                progress_value = 50 + ((i + 1) / len(pages[1:])) * 25  # Second quarter of progress
                progress_callback(int(progress_value))
            composer.append(doc)
        
        now = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        outname = f"{now}_inventory_slips.docx"
        outpath = os.path.join(output_dir, outname)
        
        if status_callback:
            status_callback("Saving document...")
        
        master.save(outpath)
        
        if status_callback:
            status_callback("Adjusting formatting...")
        
        adjust_table_font_sizes(outpath)
        
        if progress_callback:
            progress_callback(100)  # Complete progress
        
        if status_callback:
            status_callback(f"Saved to: {outpath}")
        
        # Open file if configured
        auto_open = config['SETTINGS'].getboolean('auto_open', True)
        if auto_open:
            open_file(outpath)
        
        return True, outpath
    
    except Exception as e:
        if status_callback:
            status_callback(f"Error: {e}")
        return False, str(e) 