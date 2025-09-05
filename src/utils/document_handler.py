import logging
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import re

logger = logging.getLogger(__name__)

class DocumentHandler:
    def __init__(self):
        self.doc = None
        
    def create_document(self, template_path):
        """Create document from template"""
        if not os.path.exists(template_path):
            raise ValueError(f"Template not found: {template_path}")
        self.doc = Document(template_path)
        return self.doc

    def add_content_to_table(self, records):
        """Add content to document replacing placeholders (optimized)
        NOTE: To enforce exact table/cell dimensions, set all table and cell sizes to fixed values in the DOCX template itself.
        Optionally, you can enforce cell margins in code (see below)."""
        if not records or not isinstance(records, list):
            return False
    # Optionally, enforce cell margins for all tables/cells after loading template
    def enforce_cell_margins(self, margin_dxa=100):
        """
        Set all cell margins in all tables to a fixed value (in twips).
        Call this after loading your template and before saving.
        """
        if not self.doc:
            return
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    for margin in ['top', 'bottom', 'left', 'right']:
                        margin_tag = f'w:{margin}'
                        margin_elem = OxmlElement(margin_tag)
                        margin_elem.set(qn('w:w'), str(margin_dxa))
                        margin_elem.set(qn('w:type'), 'dxa')
                        tcPr.append(margin_elem)

        try:
            # Process records in chunks of 4 (for template layout)
            for chunk_index, i in enumerate(range(0, len(records), 4)):
                chunk = records[i:i + 4]

                if chunk_index > 0:
                    self.doc.add_page_break()

                # Only process paragraphs/runs that contain placeholders
                all_paragraphs = []
                for paragraph in self.doc.paragraphs:
                    all_paragraphs.append(paragraph)
                for table in self.doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                all_paragraphs.append(paragraph)

                # Build a set of all placeholders for this chunk
                placeholders = set()
                for idx in range(1, 5):
                    placeholders.update({
                        f'{{{{Label{idx}.AcceptedDate}}}}',
                        f'{{{{Label{idx}.Vendor}}}}',
                        f'{{{{Label{idx}.ProductName}}}}',
                        f'{{{{Label{idx}.Barcode}}}}',
                        f'{{{{Label{idx}.QuantityReceived}}}}'
                    })

                # Only keep paragraphs/runs that contain any placeholder
                relevant_runs = []
                for paragraph in all_paragraphs:
                    for run in paragraph.runs:
                        if any(ph in run.text for ph in placeholders):
                            relevant_runs.append(run)

                # Replace placeholders for each record in chunk
                for idx, record in enumerate(chunk, 1):
                    replacements = {
                        f'{{{{Label{idx}.AcceptedDate}}}}': record.get('Accepted Date', ''),
                        f'{{{{Label{idx}.Vendor}}}}': record.get('Vendor', 'Unknown Vendor'),
                        f'{{{{Label{idx}.ProductName}}}}': record.get('Product Name*', ''),
                        f'{{{{Label{idx}.Barcode}}}}': record.get('Barcode*', ''),
                        f'{{{{Label{idx}.QuantityReceived}}}}': str(record.get('Quantity Received*', ''))
                    }
                    for run in relevant_runs:
                        for old_text, new_text in replacements.items():
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, str(new_text))
                                run.font.name = 'Arial'
                                if old_text == f'{{{{Label{idx}.QuantityReceived}}}}':
                                    run.font.size = Pt(12)
                                else:
                                    run.font.size = Pt(11)

                # Clean up unused placeholders for this chunk
                for idx in range(len(chunk) + 1, 5):
                    empty_replacements = {
                        f'{{{{Label{idx}.AcceptedDate}}}}': '',
                        f'{{{{Label{idx}.Vendor}}}}': '',
                        f'{{{{Label{idx}.ProductName}}}}': '',
                        f'{{{{Label{idx}.Barcode}}}}': '',
                        f'{{{{Label{idx}.QuantityReceived}}}}': ''
                    }
                    for run in relevant_runs:
                        for old_text, new_text in empty_replacements.items():
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, '')

            return True

        except Exception as e:
            logger.error(f"Failed to add content: {str(e)}")
            return False

    def save_document(self, filepath):
        """Save document"""
        try:
            os.makedirs(os.path.dirname(filepath), exist_ok=True)
            self.doc.save(filepath)
            return True
        except Exception as e:
            logger.error(f"Failed to save document: {str(e)}")
            return False