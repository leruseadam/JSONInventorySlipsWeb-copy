"""
Document Generator Module
This module provides a clean, modern API for document generation using python-docx
"""

import os
import logging
from typing import List, Dict, Any, Optional, Union
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH 
from docx.enum.section import WD_ORIENT

logger = logging.getLogger(__name__)

class DocxGenerator:
    def __init__(self):
        """Initialize a new DocxGenerator instance"""
        self.doc = None
        self.logger = logging.getLogger(__name__)
    
    def create_document(self, landscape: bool = False) -> None:
        """
        Create a new document with optional landscape orientation
        
        Args:
            landscape (bool): Whether to create document in landscape orientation
        """
        self.doc = Document()
        
        if landscape:
            section = self.doc.sections[0]
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Inches(11)
            section.page_height = Inches(8.5)
            
            # Set margins
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)

    def add_heading(self, text: str, size: float = 14, bold: bool = True, alignment: str = 'center') -> None:
        """
        Add a heading to the document
        
        Args:
            text (str): The heading text
            size (float): Font size in points
            bold (bool): Whether the heading should be bold
            alignment (str): Text alignment ('left', 'center', 'right')
        """
        paragraph = self.doc.add_paragraph()
        
        # Set alignment
        if alignment == 'center':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == 'right':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)

    def create_table(self, 
                    headers: List[str], 
                    column_widths: List[float],
                    style: str = 'Table Grid') -> None:
        """
        Create a table with specified headers and column widths
        
        Args:
            headers (List[str]): List of column header text
            column_widths (List[float]): List of column widths in inches
            style (str): Table style name (default: 'Table Grid')
        """
        # Create table
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = style
        table.autofit = False
        
        # Set column widths
        for idx, width in enumerate(column_widths):
            for cell in table.columns[idx].cells:
                cell.width = Inches(width)
                
        # Add headers with formatting
        header_row = table.rows[0]
        for idx, header in enumerate(headers):
            cell = header_row.cells[idx]
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(header)
            run.bold = True
            run.font.size = Pt(11)

    def add_table_row(self, values: List[str], font_size: float = 10) -> None:
        """
        Add a row to the last table in the document
        
        Args:
            values (List[str]): List of cell values
            font_size (float): Font size in points for the row text
        """
        if not self.doc.tables:
            raise ValueError("No table exists in document")
            
        table = self.doc.tables[-1]
        row_cells = table.add_row().cells
        
        for idx, value in enumerate(values):
            if idx < len(row_cells):
                paragraph = row_cells[idx].paragraphs[0]
                run = paragraph.add_run(str(value))
                run.font.size = Pt(font_size)

    def add_page_break(self) -> None:
        """Add a page break to the document"""
        self.doc.add_page_break()

    def save(self, filepath: str) -> bool:
        """
        Save the document to disk
        
        Args:
            filepath (str): The path where the document should be saved
            
        Returns:
            bool: True if save was successful, False otherwise
        """
        try:
            # Ensure directory exists
            os.makedirs(os.path.dirname(filepath), exist_ok=True)
            
            # Save to temp file first for safety
            temp_path = filepath + '.tmp'
            self.doc.save(temp_path)
            
            # If save was successful, move to final location
            os.replace(temp_path, filepath)
            return True
            
        except Exception as e:
            self.logger.error(f"Failed to save document: {str(e)}")
            # Clean up temp file if it exists
            if os.path.exists(temp_path):
                try:
                    os.remove(temp_path)
                except:
                    pass
            return False

    def generate_inventory_slip(self, 
                              records: List[Dict[str, Any]], 
                              vendor_name: str,
                              date: str,
                              rows_per_page: int = 20) -> None:
        """
        Generate an inventory slip with the specified records
        
        Args:
            records (List[Dict]): List of inventory records
            vendor_name (str): Name of the vendor
            date (str): Date string for the slip
            rows_per_page (int): Number of rows per page before page break
        """
        # Create document in landscape mode
        self.create_document(landscape=True)
        
        # Add title and info
        self.add_heading("Order Sheet", 14)
        self.add_heading(f"Date: {date}    Vendor: {vendor_name}", 11, bold=False)
        
        # Create table
        headers = ['Product Name', 'Barcode', 'Quantity', 'Vendor', 'Accepted Date']
        widths = [4, 2, 1, 2, 2]  # Width in inches
        self.create_table(headers, widths)
        
        # Add data rows with pagination
        for idx, record in enumerate(records):
            if idx > 0 and idx % rows_per_page == 0:
                self.add_page_break()
                # Recreate table headers on new page
                self.create_table(headers, widths)
            
            row_data = [
                str(record.get('Product Name*', ''))[:100],
                str(record.get('Barcode*', ''))[:50],
                str(record.get('Quantity Received*', ''))[:10],
                str(record.get('Vendor', ''))[:50],
                str(record.get('Accepted Date', ''))[:20]
            ]
            self.add_table_row(row_data)
