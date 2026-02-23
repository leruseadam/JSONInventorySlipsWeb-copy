"""
SimpleDocumentGenerator - Creates Word documents without using templates
for more reliable document generation
"""
import logging
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
from collections import deque
import math

logger = logging.getLogger(__name__)

class SimpleDocumentGenerator:
    def __init__(self):
        self.doc = Document()
        self._setup_document()
        
    def _setup_document(self):
        """Set up initial document properties"""
        # Set larger top/bottom margins for better appearance
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1.25)
            section.bottom_margin = Inches(1.25)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
            
    def _create_table(self, rows=2, cols=2):
        """Create a table with specified dimensions and enforce exact cell/table sizes"""
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        table.autofit = False
        col_widths = [Inches(3.5), Inches(3.5)]
        row_height = Inches(2.0)
        for col_idx, width in enumerate(col_widths):
            for cell in table.columns[col_idx].cells:
                cell.width = width
        for row_idx, row in enumerate(table.rows):
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), str(int(row_height.pt * 20)))  # twips
            trHeight.set(qn('w:hRule'), 'exact')
            trPr.append(trHeight)
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                # Set cell margins
                for margin, val in [('top', 100), ('bottom', 100), ('left', 100), ('right', 100)]:
                    margin_tag = f'w:{margin}'
                    margin_elem = OxmlElement(margin_tag)
                    margin_elem.set(qn('w:w'), str(val))
                    margin_elem.set(qn('w:type'), 'dxa')
                    tcPr.append(margin_elem)
                # Add invisible padding to all cells in the first row to enforce minimum height
                if row_idx == 0 and not cell.text.strip():
                    p = cell.add_paragraph()
                    p.add_run('\u00A0')  # non-breaking space
        return table
        
    def _add_page_number(self, current_page, total_pages):
        """Add page number to footer"""
        footer = self.doc.sections[-1].footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        page_number = paragraph.add_run(f'Page {current_page} of {total_pages}')
        page_number.font.name = 'Arial'
        page_number.font.size = Pt(10)
        
    def _add_label(self, cell, data):
        """Add formatted and centered content to a cell"""
        # Product Name
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = p.add_run(data.get('ProductName', ''))
        name_run.font.name = 'Arial'
        name_run.font.size = Pt(12)
        name_run.font.bold = True

        # Barcode
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        barcode_run = p.add_run(f"Barcode: {data.get('Barcode', '')}")
        barcode_run.font.name = 'Arial'
        barcode_run.font.size = Pt(10)

        # Quantity
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        qty_run = p.add_run(f"Quantity: {data.get('QuantityReceived', '')}")
        qty_run.font.name = 'Arial'
        qty_run.font.size = Pt(10)

        # Date and Vendor
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_vendor_run = p.add_run(f"Date: {data.get('AcceptedDate', '')} | Vendor: {data.get('Vendor', '')}")
        date_vendor_run.font.name = 'Arial'
        date_vendor_run.font.size = Pt(9)
        
    def generate_document(self, records):
        """Generate document with inventory labels"""
        try:
            if not records:
                return False, "No records provided"
                
            # Calculate total pages
            items_per_page = 4
            # Map product type strings to slot numbers
            def map_type_to_slot(ptype):
                if not ptype:
                    return None
                p = str(ptype).lower()
                if 'edibl' in p or 'edible' in p:
                    return 1
                if 'bever' in p or 'drink' in p:
                    return 2
                if 'flower' in p or 'bud' in p or 'floral' in p:
                    return 3
                if 'concentrate' in p or 'wax' in p or 'shatter' in p or 'oil' in p:
                    return 4
                return None

            # Build queues per slot
            slot_queues = {i: deque() for i in range(1, items_per_page + 1)}
            misc_queue = deque()
            for rec in records:
                ptype = rec.get('Product Type*', rec.get('Inventory Type', ''))
                slot = map_type_to_slot(ptype)
                if slot and slot in slot_queues:
                    slot_queues[slot].append(rec)
                else:
                    misc_queue.append(rec)

            # Distribute misc records into empty slot queues when possible
            while misc_queue:
                placed = False
                for i in range(1, items_per_page + 1):
                    if len(slot_queues[i]) == 0:
                        slot_queues[i].append(misc_queue.popleft())
                        placed = True
                        break
                if not placed:
                    break

            max_queue_len = max((len(q) for q in slot_queues.values()), default=0)
            if misc_queue:
                extra_pages = math.ceil(len(misc_queue) / items_per_page)
                total_pages = max_queue_len + extra_pages
            else:
                total_pages = max_queue_len

            if total_pages == 0 and records:
                total_pages = 1

            current_page = 1

            # For each page, place one item per slot if available
            for page_idx in range(total_pages):
                table = self._create_table(rows=2, cols=2)
                # slot->cell mapping: 1:(0,0),2:(0,1),3:(1,0),4:(1,1)
                for slot_num in range(1, items_per_page + 1):
                    rec = None
                    if slot_queues.get(slot_num) and len(slot_queues[slot_num]) > 0:
                        rec = slot_queues[slot_num].popleft()
                    elif misc_queue:
                        rec = misc_queue.popleft()

                    row = (slot_num - 1) // 2
                    col = (slot_num - 1) % 2
                    cell = table.cell(row, col)
                    if rec:
                        self._add_label(cell, rec)
                    else:
                        # leave cell empty but keep structure
                        pass

                # Add page number
                self._add_page_number(current_page, total_pages)

                # Add page break if not last page
                if current_page < total_pages:
                    self.doc.add_section(WD_SECTION.NEW_PAGE)

                current_page += 1
                
            return True, None
            
        except Exception as e:
            logger.error(f"Error generating document: {str(e)}")
            return False, str(e)
            
    def save(self, filepath):
        """Save the document"""
        try:
            # Ensure directory exists
            os.makedirs(os.path.dirname(filepath), exist_ok=True)
            
            # Save document
            self.doc.save(filepath)
            
            # Verify the file exists and is readable
            if not os.path.exists(filepath):
                return False, "Failed to create document file"
                
            # Try to open the file to verify it's not corrupted
            test_doc = Document(filepath)
            if not test_doc.tables:
                return False, "Generated document appears to be invalid"
                
            return True, None
            
        except Exception as e:
            logger.error(f"Error saving document: {str(e)}")
            return False, str(e)
