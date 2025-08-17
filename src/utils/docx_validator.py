"""DocxValidator - Validates and repairs Word documents"""
import logging
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.shared import Pt
import os
import tempfile
import shutil

logger = logging.getLogger(__name__)

def validate_docx(filepath):
    """
    Validate that a docx file exists and can be opened.
    
    Args:
        filepath (str): Path to the docx file
        
    Returns:
        bool: True if valid, False otherwise
    """
    validator = DocxValidator()
    is_valid, _ = validator.validate_document(filepath)
    return is_valid

class DocxValidator:
    @staticmethod
    def validate_document(file_path):
        """
        Validates a DOCX document's structure and attempts repairs if needed.
        Returns (is_valid, repaired_path)
        """
        try:
            if not os.path.exists(file_path):
                logger.error(f"Document file not found: {file_path}")
                return False, None

            # Try to open the document
            doc = Document(file_path)
            
            # Log document structure
            logger.info(f"Validating document: {file_path}")
            logger.info(f"Number of sections: {len(doc.sections)}")
            logger.info(f"Number of paragraphs: {len(doc.paragraphs)}")
            logger.info(f"Number of tables: {len(doc.tables)}")
            
            # More lenient structure validation for templates
            is_template = "template" in str(file_path).lower() or (
                hasattr(doc, 'core_properties') and 
                doc.core_properties.title and 
                "template" in str(doc.core_properties.title).lower()
            )
            
            # Basic section check
            if not doc.sections:
                if is_template:
                    logger.warning("Template has no sections, attempting to add default section")
                    doc.add_section()
                else:
                    logger.error("Document has no sections")
                    return False, None
                
            # Check for basic content structure (more lenient)
            if not doc.sections:
                logger.error("Document has no sections")
                return False, None
                
            # Basic content check - at least one paragraph or table
            if not doc.paragraphs and not doc.tables:
                logger.error("Document has no content elements")
                return False, None
                
            # Allow empty initial document for template
            if any("template" in str(doc.core_properties.title).lower() or
                  "template" in str(file_path).lower()):
                logger.info("Document appears to be a template, skipping text content check")
                return True, file_path
                
            # Less strict text content check
            has_structure = False
            # Check paragraphs
            for paragraph in doc.paragraphs:
                if paragraph._element is not None:  # Has valid XML structure
                    has_structure = True
                    break
                    
            # Check tables
            if not has_structure:
                for table in doc.tables:
                    if table._tbl is not None:  # Has valid table structure
                        has_structure = True
                        break
                        
            if not has_structure:
                logger.error("Document contains no valid content structure")
                return False, None
                
            # Validate tables
            for table in doc.tables:
                if not table._tbl:
                    logger.error("Found invalid table structure")
                    return False, None
                    
            # Check sections
            for section in doc.sections:
                if not section._sectPr:
                    logger.error("Found invalid section structure")
                    return False, None
                    
            # If everything looks good
            return True, file_path
            
        except PackageNotFoundError:
            logger.error("Invalid document package structure")
            return False, None
        except Exception as e:
            logger.error(f"Document validation error: {str(e)}")
            return False, None
            
    @staticmethod
    def repair_document(file_path):
        """
        Attempts to repair a corrupted DOCX document.
        Returns (success, repaired_path)
        """
        try:
            # Create a temporary file
            temp_dir = tempfile.mkdtemp()
            temp_path = os.path.join(temp_dir, "repaired.docx")
            
            # Try to open and resave the document
            doc = Document(file_path)
            
            # Force recalculation of all properties
            for section in doc.sections:
                section._sectPr.clear_content()
                
            # Reset font properties
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(run.font.size.pt if run.font.size else 11)
                    
            # Rebuild tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if not cell._tc:
                            continue
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(run.font.size.pt if run.font.size else 11)
                                
            # Save repaired document
            doc.save(temp_path)
            
            # Validate the repaired document
            is_valid, _ = DocxValidator.validate_document(temp_path)
            if is_valid:
                # Replace original with repaired version
                shutil.copy2(temp_path, file_path)
                return True, file_path
                
            return False, None
            
        except Exception as e:
            logger.error(f"Document repair error: {str(e)}")
            return False, None
        finally:
            # Clean up temporary files
            try:
                shutil.rmtree(temp_dir)
            except:
                pass
