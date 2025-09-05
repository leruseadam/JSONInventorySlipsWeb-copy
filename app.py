"""
Flask Inventory Slip Generator - Web application for generating inventory slips
from CSV and JSON data with support for Bamboo and Cultivera formats.
"""


"""
Flask Inventory Slip Generator - Web application for generating inventory slips
from CSV and JSON data with support for Bamboo and Cultivera formats.
"""


from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, session, send_file, send_from_directory
# Standard library imports
import os
import sys
import json
import socket
import ssl
import base64
import hmac
import hashlib
import logging
import threading
import tempfile
import urllib.request
import urllib.error
import uuid
import re
import webbrowser
import time
from functools import wraps
from io import BytesIO
import zlib
from pathlib import Path
from datetime import datetime
import os
import sys
import json
import socket
import ssl
import base64
import hmac
import hashlib
import logging
import threading
import tempfile
import urllib.request
import urllib.error
import uuid
import re
import webbrowser
import time
from functools import wraps
from io import BytesIO
import zlib
from pathlib import Path
from datetime import datetime

# Third-party imports
from flask import (
    Flask, 
    render_template, 
    request, 
    redirect, 
    url_for, 
    flash, 
    jsonify, 
    session, 
    send_file, 
    send_from_directory
)
import requests
import pandas as pd
from docxtpl import DocxTemplate
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.shared import Pt, Inches
from docxcompose.composer import Composer
import configparser
from werkzeug.utils import secure_filename
import pytesseract
from pdf2image import convert_from_path
from PIL import Image
import sqlite3

# Local imports
from src.utils.document_handler import DocumentHandler
from src.ui.app import InventorySlipGenerator

# Configure logging (must be before any logger usage)
logging.basicConfig(level=logging.INFO, 
                    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Tesseract diagnostics (must be after logger setup)
import subprocess
pytesseract.pytesseract.tesseract_cmd = "/usr/local/bin/tesseract"
try:
    version = subprocess.check_output([pytesseract.pytesseract.tesseract_cmd, "--version"], text=True)
    logger.info(f"Tesseract version: {version.strip()}")
except Exception as e:
    logger.error(f"Could not get Tesseract version: {e}")
import requests
import pandas as pd
from docxtpl import DocxTemplate
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.shared import Pt, Inches
from docxcompose.composer import Composer
import configparser
from werkzeug.utils import secure_filename
import pytesseract
from pdf2image import convert_from_path
from PIL import Image
import sqlite3

# Local imports
from src.utils.document_handler import DocumentHandler
from src.ui.app import InventorySlipGenerator
import subprocess

# Explicitly set Tesseract binary path
pytesseract.pytesseract.tesseract_cmd = "/usr/local/bin/tesseract"

# Log Tesseract version at startup for diagnostics
try:
    version = subprocess.check_output([pytesseract.pytesseract.tesseract_cmd, "--version"], text=True)
    logger.info(f"Tesseract version: {version.strip()}")
except Exception as e:
    logger.error(f"Could not get Tesseract version: {e}")


# Update the compression constants
MAX_CHUNK_SIZE = 5000  # Increased to allow larger chunks
MAX_TOTAL_SIZE = 20000  # Increased to allow larger total size
COMPRESSION_LEVEL = 9  # Maximum compression

def compress_session_data(data):
    """Compress data with improved compression and size checks"""
    try:
        # Convert DataFrame to minimal JSON if needed
        if isinstance(data, pd.DataFrame):
            data = data.to_json(orient='records', date_format='iso')
        elif not isinstance(data, str):
            data = json.dumps(data, separators=(',', ':'))

        # First level compression
        compressed = zlib.compress(data.encode('utf-8'), level=COMPRESSION_LEVEL)
        
        # If still too large, reduce data
        if len(compressed) > MAX_TOTAL_SIZE:
            if isinstance(data, str):
                try:
                    # Parse JSON to reduce content
                    parsed = json.loads(data)
                    if isinstance(parsed, list):
                        # Keep only essential fields and limit records
                        reduced = []
                        for item in parsed[:25]:  # Limit to 25 records
                            reduced_item = {k: str(v)[:50] for k, v in item.items()}  # Truncate values
                            reduced.append(reduced_item)
                        data = json.dumps(reduced, separators=(',', ':'))
                    elif isinstance(parsed, dict):
                        # Reduce dictionary size
                        data = json.dumps({k: str(v)[:50] for k, v in parsed.items()})
                except:
                    # If JSON parsing fails, truncate string
                    data = data[:1000] + "...[truncated]"
            
            # Compress reduced data
            compressed = zlib.compress(data.encode('utf-8'), level=COMPRESSION_LEVEL)

        return base64.b64encode(compressed).decode('utf-8')
    except Exception as e:
        logger.error(f"Compression error: {str(e)}")
        # Flask must be imported as the very first line
        from flask import (
            Flask, 
            render_template, 
            request, 
            redirect, 
            url_for, 
            flash, 
            jsonify, 
            session, 
            send_file, 
            send_from_directory
        )
        compressed = compress_session_data(data)
        clear_chunked_data(key)
        chunks = [compressed[i:i + MAX_CHUNK_SIZE] for i in range(0, len(compressed), MAX_CHUNK_SIZE)]
        if len(chunks) > 20:
            raise ValueError(f"Data too large: {len(chunks)} chunks needed")
        session[f'{key}_chunks'] = len(chunks)
        for i, chunk in enumerate(chunks):
            chunk_key = f'{key}_chunk_{i}'
            if len(chunk) > MAX_CHUNK_SIZE:
                raise ValueError(f"Chunk {i} exceeds maximum size")
            session[chunk_key] = chunk
        logger.info(f"Stored {len(chunks)} chunks for {key} (total size: {len(compressed)})")
        return True
    except Exception as e:
        logger.error(f"Error storing chunked data: {str(e)}")
        clear_chunked_data(key)
        return False

def get_chunked_data(key):
    """Retrieve chunked data with improved error handling"""
    from src.utils.session_storage import get_data
    session_id = session.get('session_id')
    filepath_key = key + '_filepath'
    if filepath_key in session:
        return get_data(session[filepath_key])
    try:
        num_chunks = session.get(f'{key}_chunks')
        if num_chunks is None:
            return None
        if num_chunks > 20:
            logger.error(f"Too many chunks for {key}: {num_chunks}")
            return None
        chunks = []
        for i in range(num_chunks):
            chunk = session.get(f'{key}_chunk_{i}')
            if chunk is None or len(chunk) > MAX_CHUNK_SIZE:
                logger.error(f"Invalid chunk {i} for {key}")
                return None
            chunks.append(chunk)
        encoded_data = ''.join(chunks)
        try:
            compressed = base64.b64decode(encoded_data)
            decompressed = zlib.decompress(compressed)
            return decompressed.decode('utf-8')
        except Exception as e:
            logger.error(f"Decompression error: {str(e)}")
            return None
    except Exception as e:
        logger.error(f"Error retrieving chunked data: {str(e)}")
        return None


# Configure logging
logging.basicConfig(level=logging.INFO, 
                    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)



# Configure logging
logging.basicConfig(level=logging.INFO, 
                    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Explicitly set Tesseract binary path and log diagnostics
import pytesseract
import subprocess
pytesseract.pytesseract.tesseract_cmd = "/usr/local/bin/tesseract"
try:
    version = subprocess.check_output([pytesseract.pytesseract.tesseract_cmd, "--version"], text=True)
    logger.info(f"Tesseract version: {version.strip()}")
except Exception as e:
    logger.error(f"Could not get Tesseract version: {e}")

# Constants
CONFIG_FILE = os.path.expanduser("~/inventory_generator_config.ini")

def get_downloads_dir():
    """Get the default Downloads directory for both Windows and Mac"""
    try:
        if sys.platform == "win32":
            # First try Windows known folder path
            import winreg
            from ctypes import windll, wintypes
            CSIDL_PERSONAL = 5  # Documents
            SHGFP_TYPE_CURRENT = 0  # Get current path
            buf = wintypes.create_unicode_buffer(wintypes.MAX_PATH)
            windll.shell32.SHGetFolderPathW(None, CSIDL_PERSONAL, None, SHGFP_TYPE_CURRENT, buf)
            documents = buf.value
            
            # Try registry next
            try:
                with winreg.OpenKey(winreg.HKEY_CURRENT_USER, 
                    r"SOFTWARE\Microsoft\Windows\CurrentVersion\Explorer\Shell Folders") as key:
                    downloads = winreg.QueryValueEx(key, "{374DE290-123F-4565-9164-39C4925E467B}")[0]
                return downloads
            except:
                # Fall back to Documents\Downloads
                return os.path.join(documents, "Downloads")
        else:
            # macOS and Linux
            return os.path.join(os.path.expanduser("~"), "Downloads")
    except:
        # Ultimate fallback - user's home directory
        return os.path.expanduser("~")

# Update the constants
DEFAULT_SAVE_DIR = get_downloads_dir()
UPLOAD_FOLDER = os.path.join(tempfile.gettempdir(), "inventory_generator", "uploads")

# Ensure directories exist with proper permissions
try:
    os.makedirs(DEFAULT_SAVE_DIR, exist_ok=True)
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
except Exception as e:
    logger.error(f"Error creating directories: {str(e)}")
    # Fall back to temp directory if needed
    if not os.path.exists(DEFAULT_SAVE_DIR):
        DEFAULT_SAVE_DIR = tempfile.gettempdir()

APP_VERSION = "2.0.0"
ALLOWED_EXTENSIONS = {'csv', 'json', 'docx'}
MAX_CONTENT_LENGTH = 16 * 1024 * 1024  # 16 MB max upload size

# Add new constants for API configuration
API_CONFIGS = {
    'bamboo': {
        'base_url': 'https://api-trace.getbamboo.com/shared/manifests',
        'version': 'v1',
        'auth_type': 'bearer'
    },
    'cultivera': {
        'base_url': 'https://api.cultivera.com/api',
        'version': 'v1',
        'auth_type': 'basic'
    },
    'growflow': {
        'base_url': 'https://api.growflow.com',
        'version': 'v2',
        'auth_type': 'oauth2'
    }
}


# Flask app initialization (must come before route definitions)
app = Flask(__name__,
    static_url_path='',
    static_folder='static',
    template_folder='templates'
)
app.secret_key = 'your-fixed-development-secret-key'
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 32 * 1024 * 1024  # 32MB

# PDF upload DB setup
PDF_DB_PATH = 'pdf_inventory.db'
def init_pdf_db():
    conn = sqlite3.connect(PDF_DB_PATH)
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS pdf_inventory (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        filename TEXT NOT NULL,
        upload_date TEXT NOT NULL,
        ocr_text TEXT
    )''')
    conn.commit()
    conn.close()

def save_pdf_metadata(filename, ocr_text):
    conn = sqlite3.connect(PDF_DB_PATH)
    c = conn.cursor()
    c.execute('INSERT INTO pdf_inventory (filename, upload_date, ocr_text) VALUES (?, ?, ?)',
              (filename, datetime.now().isoformat(), ocr_text))
    conn.commit()
    conn.close()

def extract_text_from_pdf(pdf_path):
    try:
        images = convert_from_path(pdf_path)
        text = ""
        for i, img in enumerate(images):
            # Preprocess image: convert to grayscale and increase contrast
            img = img.convert('L')  # Grayscale
            from PIL import ImageEnhance
            enhancer = ImageEnhance.Contrast(img)
            img = enhancer.enhance(2.0)  # Increase contrast
            logger.info(f"OCR page {i+1}: mode={img.mode}, size={img.size}")
            ocr_result = pytesseract.image_to_string(img)
            logger.info(f"OCR output for page {i+1}: {repr(ocr_result)}")
            text += ocr_result + "\n"
        return text.strip()
    except Exception as e:
        import traceback, os
        error_details = traceback.format_exc()
        tesseract_cmd = getattr(pytesseract.pytesseract, 'tesseract_cmd', None)
        env_path = os.environ.get('PATH', '')
        logger.error(f"OCR error: {e}\nDetails: {error_details}\nTesseract cmd: {tesseract_cmd}\nPATH: {env_path}")
        return f"OCR error: {e}\nDetails: {error_details}\nTesseract cmd: {tesseract_cmd}\nPATH: {env_path}"

@app.route('/upload_pdfs', methods=['GET', 'POST'])
def upload_pdfs():
    init_pdf_db()
    if request.method == 'POST':
        if 'pdfs' not in request.files:
            flash('No file part')
            return redirect(request.url)
        files = request.files.getlist('pdfs')
        saved_files = []
        os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
        for file in files:
            if file and file.filename.lower().endswith('.pdf'):
                temp_filename = file.filename
                temp_save_path = os.path.join(app.config['UPLOAD_FOLDER'], temp_filename)
                file.save(temp_save_path)
                # OCR processing
                ocr_text = extract_text_from_pdf(temp_save_path)
                # Extract Product Name from OCR text
                import re
                product_match = re.search(r'Medically Compliant\s*-\s*(.*?)\s*-\s*', ocr_text)
                product_name = product_match.group(1).strip() if product_match else 'UnknownProduct'
                # Clean product name for filename
                safe_product_name = re.sub(r'[^A-Za-z0-9_\-]', '_', product_name)[:40]
                new_filename = f"{safe_product_name}.pdf"
                new_save_path = os.path.join(app.config['UPLOAD_FOLDER'], new_filename)
                os.rename(temp_save_path, new_save_path)
                save_pdf_metadata(new_filename, ocr_text)
                saved_files.append(new_filename)
        flash(f'Uploaded: {", ".join(saved_files)}')
        return redirect(url_for('upload_pdfs'))
    return render_template('upload_pdfs.html')

@app.route('/list_pdfs')
def list_pdfs():
    init_pdf_db()
    conn = sqlite3.connect(PDF_DB_PATH)
    c = conn.cursor()
    c.execute('SELECT filename, upload_date, ocr_text FROM pdf_inventory ORDER BY upload_date DESC')
    pdfs_raw = c.fetchall()
    conn.close()
    # Parse OCR text into structured columns
    def parse_slip(text):
        import re
        # Only keep lines with expected fields
        lines = text.splitlines()
        filtered = []
        for line in lines:
            if re.search(r'(\d{4}-\d{2}-\d{2}|JSM LLC|Only B\'s|Dank Czar|Flavour Bar|Omega Distillate|Medically Compliant|SKU:|Initial Qty Issued:|Qty Received:)', line):
                filtered.append(line.strip())
        filtered_text = '\n'.join(filtered)
        # Extract fields using regex (simple version)
        date = re.search(r'\d{4}-\d{2}-\d{2}', filtered_text)
        vendor = re.search(r'JSM LLC|Only B\'s|Dank Czar|Flavour Bar|Omega Distillate', filtered_text)
        product = re.search(r'(Medically Compliant.*?)(SKU:|$)', filtered_text, re.DOTALL)
        sku = re.search(r'SKU:\s*(\d+)', filtered_text)
        qty_issued = re.search(r'Initial Qty Issued:\s*\|?\s*(\d+)?', filtered_text)
        qty_received = re.search(r'Qty Received:\s*\|?\s*(\d+)?', filtered_text)
        return {
            'date': date.group(0) if date else '',
            'vendor': vendor.group(0) if vendor else '',
            'product': product.group(1).replace('\n', ' ').strip() if product else '',
            'sku': sku.group(1) if sku else '',
            'qty_issued': qty_issued.group(1) if qty_issued else '',
            'qty_received': qty_received.group(1) if qty_received else ''
        }
    pdfs = []
    for filename, upload_date, ocr_text in pdfs_raw:
        slip = parse_slip(ocr_text or '')
        pdfs.append((filename, upload_date, slip['date'], slip['vendor'], slip['product'], slip['sku'], slip['qty_issued'], slip['qty_received']))
    return render_template('list_pdfs.html', pdfs=pdfs)

# Add security headers and session configuration for Chrome compatibility
@app.after_request
def add_security_headers(response):
    """Add security headers to make the app work better with Chrome"""
    # Remove any existing problematic headers
    response.headers.pop('X-Frame-Options', None)
    response.headers.pop('X-Content-Type-Options', None)
    
    # Add Chrome-compatible headers
    response.headers['X-Content-Type-Options'] = 'nosniff'
    response.headers['X-Frame-Options'] = 'SAMEORIGIN'
    response.headers['X-XSS-Protection'] = '1; mode=block'
    
    # Add CORS headers for local development
    response.headers['Access-Control-Allow-Origin'] = 'http://localhost:*'
    response.headers['Access-Control-Allow-Methods'] = 'GET, POST, PUT, DELETE, OPTIONS'
    response.headers['Access-Control-Allow-Headers'] = 'Content-Type, Authorization'
    response.headers['Access-Control-Allow-Credentials'] = 'true'
    
    # Add cache control headers
    response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    response.headers['Pragma'] = 'no-cache'
    response.headers['Expires'] = '0'
    
    return response

# Configure session for Chrome compatibility
app.config.update(
    SESSION_COOKIE_SECURE=False,  # Set to True only if using HTTPS
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE='Lax',  # More permissive than 'Strict' for Chrome
    SESSION_COOKIE_PATH='/',
    PERMANENT_SESSION_LIFETIME=3600,  # 1 hour
)

# Helper function to get resource path (for templates)
def resource_path(relative_path):
    """Get absolute path to resource, works for dev and for PyInstaller"""
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, relative_path)
    return os.path.join(os.path.abspath("."), relative_path)

# Load configurations or create default
def load_config():
    config = configparser.ConfigParser()
    
    # Default configurations
    config['PATHS'] = {
        'template_path': os.path.join(os.path.dirname(__file__), "templates/documents/InventorySlips.docx"),
        'output_dir': DEFAULT_SAVE_DIR,  # Use the new DEFAULT_SAVE_DIR
        'recent_files': '',
        'recent_urls': ''
    }
    
    config['SETTINGS'] = {
        'items_per_page': '4',
        'auto_open': 'true',
        'theme': 'dark',
        'font_size': '12'
    }
    
    # Load existing config if it exists
    if os.path.exists(CONFIG_FILE):
        config.read(CONFIG_FILE)
    else:
        # Create config file with defaults
        with open(CONFIG_FILE, 'w') as f:
            config.write(f)
    
    return config

def save_config(config):
    with open(CONFIG_FILE, 'w') as f:
        config.write(f)

# Helper to adjust font sizes after rendering
def adjust_table_font_sizes(doc_path):
    """
    Post-process a DOCX file to dynamically adjust font size inside table cells based on thresholds.
    """
    thresholds = [
        (30, 12),   # <=30 chars → 12pt
        (45, 10),   # <=45 chars → 10pt
        (60, 8),    # <=60 chars → 8pt
        (float('inf'), 7)  # >60 chars → 7pt
    ]

    def get_font_size(text_len):
        for limit, size in thresholds:
            if text_len <= limit:
                return size
        return 7  # Fallback

    try:
        doc = Document(doc_path)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text.strip()
                        if not text:
                            continue

                        # If line is Product Name (first line), force 10pt
                        if paragraph == cell.paragraphs[0]:
                            font_size = 10
                        else:
                            font_size = get_font_size(len(text))

                        for run in paragraph.runs:
                            try:
                                run.font.size = Pt(font_size)
                            except Exception as e:
                                logger.warning(f"Could not set font size for run: {e}")
                                continue

        # Save to a temporary file first
        temp_path = doc_path + ".font_adjust.tmp"
        doc.save(temp_path)
        
        # Validate the document after font adjustments
        if validate_docx(temp_path):
            import shutil
            shutil.move(temp_path, doc_path)
        else:
            logger.warning("Document validation failed after font adjustment, keeping original")
            if os.path.exists(temp_path):
                os.remove(temp_path)
                
    except Exception as e:
        logger.error(f"Error adjusting font sizes: {str(e)}")
        # If font adjustment fails, the document should still be usable
        if os.path.exists(temp_path):
            os.remove(temp_path)

# Open files after saving
def open_file(path):
    """Open files using the default system application"""
    try:
        if sys.platform == "win32":
            os.startfile(path)
        elif sys.platform == "darwin":  # macOS
            os.system(f'open "{path}"')
        else:  # linux variants
            os.system(f'xdg-open "{path}"')
    except Exception as e:
        logger.error(f"Error opening file: {e}")
        flash(f"Error opening file: {e}", "error")

# Split records into chunks
def chunk_records(records, chunk_size=4):
    for i in range(0, len(records), chunk_size):
        yield records[i:i + chunk_size]

# Check if file extension is allowed
def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Process and save inventory slips
def run_full_process_inventory_slips(selected_df, config, status_callback=None, progress_callback=None):
    if selected_df.empty:
        if status_callback:
            status_callback("Error: No data selected.")
        return False, "No data selected."

def run_full_process_inventory_slips(selected_df, config, status_callback=None, progress_callback=None):
    # ...existing code...
    
    try:
        # Get vendor name from first row
        vendor_name = selected_df['Vendor'].iloc[0] if not selected_df.empty else "Unknown"
        # Clean vendor name (remove special characters and spaces)
        vendor_name = "".join(c for c in vendor_name if c.isalnum() or c.isspace()).strip()
        # Get today's date
        today_date = datetime.now().strftime("%Y%m%d")
        # Create filename
        outname = f"{today_date}_{vendor_name}_Slips.docx"
        output_dir = config['PATHS']['output_dir']
        if not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        outpath = os.path.join(output_dir, outname)
        
        # ...rest of existing code...
        
        # Get settings from config
        items_per_page = int(config['SETTINGS'].get('items_per_page', '4'))
        template_path = config['PATHS'].get('template_path')
        if not template_path or not os.path.exists(template_path):
            template_path = os.path.join(os.path.dirname(__file__), "templates/documents/InventorySlips.docx")
            if not os.path.exists(template_path):
                raise ValueError(f"Template file not found at: {template_path}")
        
        if status_callback:
            status_callback("Processing data...")

        # Clean and validate the data
        records = selected_df.to_dict(orient="records")
        cleaned_records = []
        
        for record in records:
            cleaned_record = {}
            for key, value in record.items():
                # Convert all values to strings and clean them
                if value is None:
                    cleaned_record[key] = ""
                else:
                    # Remove any problematic characters and limit length
                    cleaned_value = str(value).strip()
                    # Remove any non-printable characters except newlines and tabs
                    cleaned_value = ''.join(char for char in cleaned_value if char.isprintable() or char in '\n\t')
                    # Limit length to prevent overflow
                    cleaned_record[key] = cleaned_value[:200] if len(cleaned_value) > 200 else cleaned_value
            cleaned_records.append(cleaned_record)
        
        pages = []

        # Process records in chunks of 4 (or configured size)
        total_chunks = (len(cleaned_records) + items_per_page - 1) // items_per_page
        current_chunk = 0

        for chunk in chunk_records(cleaned_records, items_per_page):
            current_chunk += 1
            if progress_callback:
                progress = (current_chunk / total_chunks) * 50
                progress_callback(int(progress))

            if status_callback:
                status_callback(f"Generating page {current_chunk} of {total_chunks}...")

            try:
                # Create a fresh template instance for each chunk
                tpl = DocxTemplate(template_path)
                context = {}

                # Fill context with records - modified vendor handling
                for idx, record in enumerate(chunk, 1):
                    # Get vendor info, using full vendor name if available
                    vendor_name = record.get("Vendor", "")
                    # If vendor is in format "license - name", extract just the name
                    if " - " in vendor_name:
                        vendor_name = vendor_name.split(" - ")[1]
                    
                    # Ensure all values are strings and not too long
                    context[f"Label{idx}"] = {
                        "ProductName": str(record.get("Product Name*", ""))[:100],
                        "Barcode": str(record.get("Barcode*", ""))[:50],
                        "AcceptedDate": str(record.get("Accepted Date", ""))[:20],
                        "QuantityReceived": str(record.get("Quantity Received*", ""))[:20],
                        "Vendor": str(vendor_name or "Unknown Vendor")[:50],
                        "ProductType": str(record.get("Product Type*", ""))[:50]
                    }

                # Fill remaining slots with empty values
                for i in range(len(chunk) + 1, items_per_page + 1):
                    context[f"Label{i}"] = {
                        "ProductName": "",
                        "Barcode": "",
                        "AcceptedDate": "",
                        "QuantityReceived": "",
                        "Vendor": "",
                        "ProductType": ""
                    }

                # Render template with context
                tpl.render(context)
                
                # Save to BytesIO with proper error handling
                output = BytesIO()
                tpl.save(output)
                output.seek(0)
                
                # Create document from BytesIO
                doc = Document(output)
                pages.append(doc)

            except Exception as e:
                logger.error(f"Error generating page {current_chunk}: {e}")
                raise ValueError(f"Error generating page {current_chunk}: {e}")

        if not pages:
            return False, "No documents generated."

        # Combine pages with better error handling
        if status_callback:
            status_callback("Combining pages...")

        try:
            master = pages[0]
            composer = Composer(master)
            
            for i, doc in enumerate(pages[1:]):
                if progress_callback:
                    progress = 50 + ((i + 1) / len(pages[1:])) * 40
                    progress_callback(int(progress))
                
                # Add page break before appending
                if hasattr(composer, 'doc') and composer.doc.paragraphs:
                    composer.doc.paragraphs[-1].add_run().add_break()
                
                composer.append(doc)

            # Save final document with proper error handling
            now = datetime.now().strftime("%Y%m%d_%H%M%S")
            outname = f"inventory_slips_{now}.docx"
            outpath = os.path.join(config['PATHS']['output_dir'], outname)

            if status_callback:
                status_callback("Saving document...")

            # Save to a temporary file first, then move to final location
            temp_path = outpath + ".tmp"
            master.save(temp_path)
            
            # Validate the saved document
            if not validate_docx(temp_path):
                raise ValueError("Generated document is corrupted")
            
            # Move to final location
            import shutil
            shutil.move(temp_path, outpath)

            # Add page numbers to footer
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            def add_page_number(footer):
                paragraph = footer.paragraphs[0]
                run = paragraph.add_run()
                fldChar1 = OxmlElement('w:fldChar')
                fldChar1.set(qn('w:fldCharType'), 'begin')
                instrText = OxmlElement('w:instrText')
                instrText.text = 'PAGE'
                fldChar2 = OxmlElement('w:fldChar')
                fldChar2.set(qn('w:fldCharType'), 'end')
                run._r.append(fldChar1)
                run._r.append(instrText)
                run._r.append(fldChar2)
            for section in master.sections:
                add_page_number(section.footer)
            master.save(temp_path)
            # Adjust font sizes
            if status_callback:
                status_callback("Adjusting formatting...")
            adjust_table_font_sizes(outpath)
            if progress_callback:
                progress_callback(100)
            return True, outpath

        except Exception as e:
            logger.error(f"Error combining or saving documents: {e}")
            raise ValueError(f"Error combining or saving documents: {e}")

    except Exception as e:
        if status_callback:
            status_callback(f"Error: {str(e)}")
        logger.error(f"Error in run_full_process_inventory_slips: {str(e)}")
        return False, str(e)

# Parse Bamboo transfer schema JSON
def parse_bamboo_data(json_data):
    if not json_data:
        return pd.DataFrame()
    
    try:
        # Get vendor information
        from_license_number = json_data.get("from_license_number", "")
        from_license_name = json_data.get("from_license_name", "")
        vendor_meta = f"{from_license_number} - {from_license_name}"
        
        # Get transfer date
        raw_date = json_data.get("est_arrival_at", "") or json_data.get("transferred_at", "")
        accepted_date = raw_date.split("T")[0] if "T" in raw_date else raw_date
        
        # Process inventory items
        items = json_data.get("inventory_transfer_items", [])
        logger.info(f"Bamboo data: found {len(items)} inventory_transfer_items")
        records = []
        
        for item in items:
            # Extract THC and CBD content from lab_result_data if available
            thc_content = ""
            cbd_content = ""
            
            lab_data = item.get("lab_result_data", {})
            if lab_data and "potency" in lab_data:
                for potency_item in lab_data["potency"]:
                    if potency_item.get("type") == "total-thc":
                        thc_content = f"{potency_item.get('value', '')}%"
                    elif potency_item.get("type") == "total-cbd":
                        cbd_content = f"{potency_item.get('value', '')}%"
            
            records.append({
                "Product Name*": item.get("product_name", ""),
                "Product Type*": item.get("inventory_type", ""),
                "Quantity Received*": item.get("qty", ""),
                "Barcode*": item.get("inventory_id", "") or item.get("external_id", ""),
                "Accepted Date": accepted_date,
                "Vendor": vendor_meta,
                "Strain Name": item.get("strain_name", ""),
                "THC Content": thc_content,
                "CBD Content": cbd_content,
                "Source System": "Bamboo"
            })
        
        return pd.DataFrame(records)
    
    except Exception as e:
        raise ValueError(f"Failed to parse Bamboo transfer data: {e}")

# Parse Cultivera JSON
def parse_cultivera_data(json_data):
    if not json_data:
        return pd.DataFrame()
    
    try:
        # Check if Cultivera format
        if not json_data.get("data") or not isinstance(json_data.get("data"), dict):
            raise ValueError("Not a valid Cultivera format")
        
        data = json_data.get("data", {})
        manifest = data.get("manifest", {})
        
        # Get vendor information
        from_license = manifest.get("from_license", {})
        vendor_name = from_license.get("name", "")
        vendor_license = from_license.get("license_number", "")
        vendor_meta = f"{vendor_license} - {vendor_name}" if vendor_license and vendor_name else "Unknown Vendor"
        
        # Get transfer date
        created_at = manifest.get("created_at", "")
        accepted_date = created_at.split("T")[0] if "T" in created_at else created_at
        
        # Process inventory items
        items = manifest.get("items", [])
        records = []
        
        for item in items:
            # Extract product info
            product = item.get("product", {})
            
            # Extract THC and CBD content
            thc_content = ""
            cbd_content = ""
            
            test_results = item.get("test_results", [])
            if test_results:
                for result in test_results:
                    if "thc" in result.get("type", "").lower():
                        thc_content = f"{result.get('percentage', '')}%"
                    elif "cbd" in result.get("type", "").lower():
                        cbd_content = f"{result.get('percentage', '')}%"
            
            records.append({
                "Product Name*": product.get("name", ""),
                "Product Type*": product.get("category", ""),
                "Quantity Received*": item.get("quantity", ""),
                "Barcode*": item.get("barcode", "") or item.get("id", ""),
                "Accepted Date": accepted_date,
                "Vendor": vendor_meta,
                "Strain Name": product.get("strain_name", ""),
                "THC Content": thc_content,
                "CBD Content": cbd_content,
                "Source System": "Cultivera"
            })
        
        return pd.DataFrame(records)
    
    except Exception as e:
        raise ValueError(f"Failed to parse Cultivera data: {e}")

def parse_growflow_data(json_data):
    """Parse GrowFlow JSON format into common fields"""
    try:
        if not ('inventory_transfer_items' in json_data and 
                'from_license_number' in json_data and 
                'from_license_name' in json_data):
            return pd.DataFrame()
        
        vendor_meta = f"{json_data.get('from_license_number', '')} - {json_data.get('from_license_name', 'Unknown Vendor')}"
        raw_date = json_data.get("est_arrival_at", "") or json_data.get("transferred_at", "")
        accepted_date = raw_date.split("T")[0] if "T" in raw_date else raw_date
        
        items = json_data.get("inventory_transfer_items", [])
        mapped_data = []
        
        for item in items:
            potency_data = item.get("lab_result_data", {}).get("potency", [])
            thc_value = next((p.get('value') for p in potency_data if p.get('type') in ["total-thc", "thc"]), 0)
            cbd_value = next((p.get('value') for p in potency_data if p.get('type') in ["total-cbd", "cbd"]), 0)
            
            mapped_item = {
                "Product Name*": item.get("product_name", ""),
                "Product Type*": item.get("inventory_type", ""),
                "Quantity Received*": item.get("qty", ""),
                "Barcode*": item.get("product_sku", "") or item.get("inventory_id", ""),
                "Accepted Date": accepted_date,
                "Vendor": vendor_meta,
                "Strain Name": item.get("strain_name", ""),
                "THC Content": f"{thc_value}%",
                "CBD Content": f"{cbd_value}%",
                "Source System": "GrowFlow"
            }
            mapped_data.append(mapped_item)
        
        return pd.DataFrame(mapped_data)
    
    except Exception as e:
        logger.error(f"Error parsing GrowFlow data: {str(e)}")
        return pd.DataFrame()

def parse_inventory_json(json_data):
    """
    Detects and parses JSON format accordingly
    Returns tuple of (DataFrame, format_type)
    """
    if not json_data:
        logger.info("No data provided to parse_inventory_json.")
        print("No data provided to parse_inventory_json.")
        return None, "No data provided"
    try:
        if isinstance(json_data, str):
            json_data = json.loads(json_data)
        # Try parsing as Bamboo
        if "inventory_transfer_items" in json_data:
            df = parse_bamboo_data(json_data)
            logger.info(f"Parsed Bamboo format, records: {len(df) if df is not None else 0}")
            print(f"Parsed Bamboo format, records: {len(df) if df is not None else 0}")
            return df, "Bamboo"
        # Try parsing as Cultivera
        elif "data" in json_data and isinstance(json_data["data"], dict) and "manifest" in json_data["data"]:
            df = parse_cultivera_data(json_data)
            logger.info(f"Parsed Cultivera format, records: {len(df) if df is not None else 0}")
            print(f"Parsed Cultivera format, records: {len(df) if df is not None else 0}")
            return df, "Cultivera"
        # Try parsing as GrowFlow
        elif "document_schema_version" in json_data:
            df = parse_growflow_data(json_data)
            logger.info(f"Parsed GrowFlow format, records: {len(df) if df is not None else 0}")
            print(f"Parsed GrowFlow format, records: {len(df) if df is not None else 0}")
            return df, "GrowFlow"
        else:
            logger.info("Unknown JSON format in parse_inventory_json.")
            print("Unknown JSON format in parse_inventory_json.")
            return None, "Unknown JSON format"
    except json.JSONDecodeError:
        logger.error("Invalid JSON data in parse_inventory_json.")
        print("Invalid JSON data in parse_inventory_json.")
        return None, "Invalid JSON data"
    except Exception as e:
        logger.error(f"Error parsing data in parse_inventory_json: {str(e)}")
        print(f"Error parsing data in parse_inventory_json: {str(e)}")
        return None, f"Error parsing data: {str(e)}"

# Process CSV data
def process_csv_data(df):
    try:
        # Strip whitespace from column names
        df.columns = [col.strip() for col in df.columns]
        logger.info(f"Original columns: {df.columns.tolist()}")
        
        # First, ensure column names are unique by adding a suffix if needed
        df.columns = [f"{col}_{i}" if df.columns.tolist().count(col) > 1 else col 
                     for i, col in enumerate(df.columns)]
        logger.info(f"Columns after ensuring uniqueness: {df.columns.tolist()}")
        
        # Map column names to expected format
        col_map = {
            "Product Name*": "Product Name*",
            "Product Name": "Product Name*",
            "Quantity Received": "Quantity Received*",
            "Quantity*": "Quantity Received*",
            "Quantity": "Quantity Received*",
            "Lot Number*": "Barcode*",
            "Barcode": "Barcode*",
            "Lot Number": "Barcode*",
            "Accepted Date": "Accepted Date",
            "Vendor": "Vendor",
            "Strain Name": "Strain Name",
            "Product Type*": "Product Type*",
            "Product Type": "Product Type*",
            "Inventory Type": "Product Type*"
        }
        
        # Now rename columns according to our mapping
        new_columns = {}
        target_counts = {}  # Keep track of how many times we've used each target name
        
        for col in df.columns:
            base_col = col.split('_')[0]  # Remove any suffix
            if base_col in col_map:
                target_name = col_map[base_col]
                # If we've seen this target name before, add a suffix
                if target_name in target_counts:
                    target_counts[target_name] += 1
                    new_columns[col] = f"{target_name}_{target_counts[target_name]}"
                else:
                    target_counts[target_name] = 0
                    new_columns[col] = target_name
            else:
                new_columns[col] = col
        
        logger.info(f"Column mapping: {new_columns}")
        df = df.rename(columns=new_columns)
        logger.info(f"Columns after renaming: {df.columns.tolist()}")
        
        # Ensure required columns exist
        required_cols = ["Product Name*", "Barcode*"]
        missing_cols = [col for col in required_cols if not any(col in c for c in df.columns)]
        
        if missing_cols:
            return None, f"CSV is missing required columns: {', '.join(missing_cols)}"
        
        # Set default values for missing columns
        if not any("Vendor" in c for c in df.columns):
            df["Vendor"] = "Unknown Vendor"
        else:
            vendor_col = next(c for c in df.columns if "Vendor" in c)
            df[vendor_col] = df[vendor_col].fillna("Unknown Vendor")
        
        if not any("Accepted Date" in c for c in df.columns):
            today = datetime.today().strftime("%Y-%m-%d")
            df["Accepted Date"] = today
        
        if not any("Product Type*" in c for c in df.columns):
            df["Product Type*"] = "Unknown"
        
        if not any("Strain Name" in c for c in df.columns):
            df["Strain Name"] = ""
        
        # Sort if possible
        try:
            sort_cols = []
            if any("Product Type*" in c for c in df.columns):
                sort_cols.append(next(c for c in df.columns if "Product Type*" in c))
            if any("Product Name*" in c for c in df.columns):
                sort_cols.append(next(c for c in df.columns if "Product Name*" in c))
            
            if sort_cols:
                df = df.sort_values(sort_cols, ascending=[True, True])
        except:
            pass  # If sorting fails, continue without sorting
        
        # Final check for duplicate columns
        if len(df.columns) != len(set(df.columns)):
            duplicates = [col for col in df.columns if df.columns.tolist().count(col) > 1]
            logger.error(f"Duplicate columns found: {duplicates}")
            return None, f"Duplicate columns found: {', '.join(duplicates)}"
        
        return df, "Success"
    
    except Exception as e:
        logger.error(f"Error in process_csv_data: {str(e)}", exc_info=True)
        return None, f"Failed to process CSV data: {e}"

def limit_dataframe_for_session(df, max_rows=25):
    """Limit DataFrame size to prevent session cookie from becoming too large"""
    if len(df) > max_rows:
        logger.warning(f"DataFrame has {len(df)} rows, limiting to {max_rows} for session storage")
        return df.head(max_rows)
    return df

def store_chunked_data(key, data):
    """Store data with improved chunking and size validation"""
    try:
        # Compress data first
        compressed = compress_session_data(data)
        
        # Clear existing chunks
        clear_chunked_data(key)
        
        # Split into smaller chunks
        chunks = [compressed[i:i + MAX_CHUNK_SIZE] for i in range(0, len(compressed), MAX_CHUNK_SIZE)]
        
        if len(chunks) > 20:  # Limit number of chunks
            raise ValueError(f"Data too large: {len(chunks)} chunks needed")
        
        # Store chunks with size validation
        session[f'{key}_chunks'] = len(chunks)
        for i, chunk in enumerate(chunks):
            chunk_key = f'{key}_chunk_{i}'
            if len(chunk) > MAX_CHUNK_SIZE:
                raise ValueError(f"Chunk {i} exceeds maximum size")
            session[chunk_key] = chunk
            
        logger.info(f"Stored {len(chunks)} chunks for {key} (total size: {len(compressed)})")
        return True
        
    except Exception as e:
        logger.error(f"Error storing chunked data: {str(e)}")
        clear_chunked_data(key)  # Clean up on error
        return False

def get_chunked_data(key):
    """Retrieve chunked data with improved error handling"""
    try:
        num_chunks = session.get(f'{key}_chunks')
        if num_chunks is None:
            return None
            
        # Validate chunk count
        if num_chunks > 20:  # Safety check
            logger.error(f"Too many chunks for {key}: {num_chunks}")
            return None
            
        # Reconstruct data
        chunks = []
        for i in range(num_chunks):
            chunk = session.get(f'{key}_chunk_{i}')
            if chunk is None or len(chunk) > MAX_CHUNK_SIZE:
                logger.error(f"Invalid chunk {i} for {key}")
                return None
            chunks.append(chunk)
            
        # Combine and decompress
        encoded_data = ''.join(chunks)
        try:
            compressed = base64.b64decode(encoded_data)
            decompressed = zlib.decompress(compressed)
            return decompressed.decode('utf-8')
        except Exception as e:
            logger.error(f"Decompression error: {str(e)}")
            return None
            
    except Exception as e:
        logger.error(f"Error retrieving chunked data: {str(e)}")
        return None

def clear_chunked_data(key):
    """Remove all chunks for a given key"""
    try:
        num_chunks = session.get(f'{key}_chunks')
        if num_chunks is not None:
            for i in range(num_chunks):
                session.pop(f'{key}_chunk_{i}', None)
        session.pop(f'{key}_chunks', None)
    except Exception as e:
        logger.error(f"Error clearing chunked data: {str(e)}")

def cleanup_temp_files():
    """Clean up any temporary files that might be left behind"""
    try:
        import glob
        temp_patterns = [
            os.path.join(DEFAULT_SAVE_DIR, "*.tmp"),
            os.path.join(DEFAULT_SAVE_DIR, "*.font_adjust.tmp"),
            os.path.join(UPLOAD_FOLDER, "*.tmp")
        ]
        
        for pattern in temp_patterns:
            for temp_file in glob.glob(pattern):
                try:
                    os.remove(temp_file)
                    logger.info(f"Cleaned up temporary file: {temp_file}")
                except Exception as e:
                    logger.warning(f"Could not remove temporary file {temp_file}: {e}")
    except Exception as e:
        logger.error(f"Error during cleanup: {e}")

def create_robust_inventory_slip(selected_df, config, status_callback=None):
    try:
        # Get vendor name
        vendor_name = selected_df['Vendor'].iloc[0] if not selected_df.empty else "Unknown"
        if " - " in vendor_name:
            vendor_name = vendor_name.split(" - ")[1]
        vendor_name = "".join(c for c in vendor_name if c.isalnum() or c.isspace()).strip()
        
        # Create filename
        today_date = datetime.now().strftime("%Y%m%d")
        outname = f"{today_date}_{vendor_name}_OrderSheet.docx"
        outpath = os.path.join(config['PATHS']['output_dir'], outname)

        # Create new document with landscape orientation first
        doc = Document()
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        
        # Set margins
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

        # Add title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("Order Sheet")
        run.bold = True
        run.font.size = Pt(14)

        # Add date and vendor info
        info = doc.add_paragraph()
        run = info.add_run(f"Date: {today_date}    Vendor: {vendor_name}")
        run.font.size = Pt(11)

        # Create table with proper dimensions
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        table.autofit = False  # Prevent autofit to keep our dimensions

        # Set column widths proportionally
        widths = [4, 2, 1, 2, 2, 2]  # Total = 13 inches
        total_width = sum(widths)
        page_width = 10  # Actual usable width after margins
        
        for i, width in enumerate(widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(width * page_width / total_width)

        # Add headers
        headers = ['Product Name', 'Barcode', 'Quantity', 'Vendor', 'Accepted Date']
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(header)
            run.bold = True
            run.font.size = Pt(11)

        # Add data rows with pagination
        rows_per_page = 20  # Adjust based on page size and margins
        current_row = 0
        
        for _, row in selected_df.iterrows():
            if current_row > 0 and current_row % rows_per_page == 0:
                doc.add_page_break()
                # Add header row in new page
                # Create table with proper dimensions
                table = doc.add_table(rows=1, cols=6)
                table.style = 'Table Grid'
                table.autofit = False  # Keep autofit false for manual width control

                # Set column widths proportionally for better fit
                widths = [5, 2, 0.75, 1.5, 1.5, 1.75]  # Adjusted widths for better proportions
                total_width = sum(widths)
                page_width = 10  # Actual usable width after margins

                # Apply widths to first table
                for i, width in enumerate(widths):
                    for cell in table.columns[i].cells:
                        cell.width = Inches(width * page_width / total_width)

                # When creating new tables for additional pages, use the same settings
                if current_row > 0 and current_row % rows_per_page == 0:
                    doc.add_page_break()
                    table = doc.add_table(rows=1, cols=6)
                    table.style = 'Table Grid'
                    table.autofit = False  # Keep autofit false for consistent width control
                    
                    # Apply same widths to new table
                    for i, width in enumerate(widths):
                        for cell in table.columns[i].cells:
                            cell.width = Inches(width * page_width / total_width)
                current_row += 1
                
                # Add headers to new page
                for i, header in enumerate(headers):
                    cell = table.cell(0, i)
                    paragraph = cell.paragraphs[0]
                    run = paragraph.add_run(header)
                    run.bold = True
                    run.font.size = Pt(11)

            row_cells = table.add_row().cells
            data = [
                str(row.get('Product Name*', ''))[:100],
                str(row.get('Barcode*', ''))[:50],
                str(row.get('Quantity Received*', ''))[:5],
                str(row.get('Vendor', ''))[:20],
                str(row.get('Accepted Date', ''))[:10]
            ]
            
            for i, value in enumerate(data):
                paragraph = row_cells[i].paragraphs[0]
                run = paragraph.add_run(value)
                run.font.size = Pt(10)
            
            current_row += 1

        # Save document
        doc.save(outpath)
        
        if os.path.exists(outpath):
            return True, outpath
            
        return False, "Failed to create document"

    except Exception as e:
        logger.error(f"Error in create_robust_inventory_slip: {str(e)}")
        if os.path.exists(outpath):
            try:
                os.remove(outpath)
            except:
                pass
        return False, str(e)

@app.route('/paste-json', methods=['POST'])
def paste_json():
    try:
        data = request.get_json()
        pasted_json = data.get('json_data', '')
        if not pasted_json:
            return jsonify({'success': False, 'message': 'No JSON data provided.'}), 400

        try:
            parsed = json.loads(pasted_json)
        except Exception as e:
            return jsonify({'success': False, 'message': f'Invalid JSON: {str(e)}'}), 400

        result_df, format_type = parse_inventory_json(parsed)
        if result_df is None or result_df.empty:
            return jsonify({'success': False, 'message': 'Could not process pasted JSON data.'}), 400

        # Store data using chunked storage
        store_chunked_data('df_json', result_df)
        # Only store raw data if it's small enough
        if len(pasted_json) < 2000:
            store_chunked_data('raw_json', pasted_json)
        else:
            store_chunked_data('raw_json', {"type": "large_json", "size": len(pasted_json)})
        session['format_type'] = format_type

        return jsonify({'success': True, 'redirect': url_for('data_view')})
    except Exception as e:
        return jsonify({'success': False, 'message': f'Error: {str(e)}'}), 500

@app.route('/upload-csv', methods=['POST'])
def upload_csv():
    if 'file' not in request.files:
        flash('No file part')
        return redirect(url_for('index'))
    file = request.files['file']
    if file.filename == '':
        flash('No selected file')
        return redirect(url_for('index'))
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        try:
            df = pd.read_csv(filepath)
            processed_df, msg = process_csv_data(df)
            if processed_df is None:
                flash(msg)
                return redirect(url_for('index'))
            
            # Store data using chunked storage
            store_chunked_data('df_json', processed_df)
            # Only store raw data if it's small enough
            raw_json = df.to_json(orient='records', default_handler=str)
            if len(raw_json) < 2000:
                store_chunked_data('raw_json', raw_json)
            else:
                store_chunked_data('raw_json', {"type": "large_csv", "rows": len(df), "columns": list(df.columns)})
            session['format_type'] = 'CSV'
            
            flash('CSV uploaded and processed successfully')
            return redirect(url_for('data_view'))
        except Exception as e:
            flash(f'Failed to process CSV: {str(e)}')
            return redirect(url_for('index'))
    else:
        flash('Invalid file type')
        return redirect(url_for('index'))

# Then, update the URL loading function
@app.route('/load-url', methods=['POST'])
def load_url():
    import traceback
    print("Entered load_url route")
    try:
        url = request.form.get('url')
        print(f"URL received: {url}")
        if not url:
            print("No URL provided")
            flash('Please enter a URL')
            return redirect(url_for('index'))
        result_df, format_type, raw_data = load_from_url(url)
        print(f"Result DataFrame: {result_df}")
        print(f"Format type: {format_type}")
        print(f"Raw data: {str(raw_data)[:500]}")
        for key in ['df_json', 'raw_json']:
            clear_chunked_data(key)
        if not store_chunked_data('df_json', result_df):
            print("Failed to store DataFrame")
            raise Exception("Failed to store DataFrame")
        print("Redirecting to data_view")
        return redirect(url_for('data_view'))
    except Exception as e:
        logger.error(f'Error loading data from URL: {str(e)}\n{traceback.format_exc()}')
        print(f"Error loading data from URL: {str(e)}\n{traceback.format_exc()}")
        flash(f'Error loading data: {str(e)}')
    return redirect(url_for('index'))
    
def handle_bamboo_url(url):
    try:
        result_df, format_type, raw_data = load_from_url(url)
        if result_df is None or result_df.empty:
            flash('Could not process Bamboo data from URL', 'error')
            return redirect(url_for('index'))
        
        store_chunked_data('df_json', result_df)
        
        # Store raw data for transfer info extraction
        if raw_data:
            store_chunked_data('raw_json', json.dumps(raw_data))
        
        session['format_type'] = format_type
        flash(f'{format_type} data loaded successfully', 'success')
        return redirect(url_for('data_view'))
    except Exception as e:
        logger.error(f'Error loading Bamboo URL: {str(e)}', exc_info=True)
        flash(f'Error loading Bamboo data: {str(e)}', 'error')
        return redirect(url_for('index'))

def load_from_url(url):
    """Download JSON or CSV data from a URL and return as DataFrame, format_type, and raw_data."""
    import traceback
    try:
        import ijson
        with requests.get(url, timeout=120, stream=True) as response:
            response.raise_for_status()
            content_type = response.headers.get('Content-Type', '').lower()
            raw_text = response.text
            print(f"Raw response (first 500 chars): {raw_text[:500]}")
            # Decide how to parse based on first non-whitespace character
            first_char = raw_text.lstrip()[0] if raw_text.lstrip() else ''
            if 'application/json' in content_type or url.lower().endswith('.json'):
                if first_char == '{':
                    # Top-level object, use json.loads
                    data = json.loads(raw_text)
                    df, format_type = parse_inventory_json(data)
                    return df, format_type, data
                elif first_char == '[':
                    # Top-level array, use ijson
                    parser = ijson.items(response.raw, 'item')
                    data = [item for item in parser]
                    df, format_type = parse_inventory_json(data)
                    return df, format_type, data
                else:
                    raise ValueError("Unknown JSON structure")
            elif 'text/csv' in content_type or url.lower().endswith('.csv'):
                df = pd.read_csv(response.raw)
                df, msg = process_csv_data(df)
                return df, 'CSV', None
            else:
                try:
                    if first_char == '{':
                        data = json.loads(raw_text)
                        df, format_type = parse_inventory_json(data)
                        return df, format_type, data
                    elif first_char == '[':
                        parser = ijson.items(response.raw, 'item')
                        data = [item for item in parser]
                        df, format_type = parse_inventory_json(data)
                        return df, format_type, data
                    else:
                        raise ValueError("Unknown JSON structure")
                except Exception:
                    try:
                        df = pd.read_csv(response.raw)
                        df, msg = process_csv_data(df)
                        return df, 'CSV', None
                    except Exception as e:
                        raise ValueError(f"Unsupported data format or failed to parse: {e}")
    except Exception as e:
        logger.error(f"Failed to load data from URL: {str(e)}\n{traceback.format_exc()}")
        print(f"Failed to load data from URL: {str(e)}\n{traceback.format_exc()}")
        raise ValueError(f"Failed to load data from URL: {e}")
    


# Update data view to handle chunked data properly
@app.route('/data-view')
def data_view():
    try:
        # Get chunked data from session
        df_json = get_chunked_data('df_json')
        format_type = session.get('format_type')

        if df_json is None:
            flash('No data available. Please load data first.')
            return redirect(url_for('index'))

        try:
            if isinstance(df_json, list):
                df = pd.DataFrame(df_json)
            else:
                df = pd.read_json(df_json, orient='records')
        except Exception as e:
            logger.error(f"Error parsing JSON data: {str(e)}")
            flash('Error loading data. Please try again.')
            return redirect(url_for('index'))
        
        # Debug logging
        logger.info(f"DataFrame shape: {df.shape}")
        logger.info(f"DataFrame columns: {df.columns.tolist()}")
        if not df.empty:
            logger.info(f"First row data: {df.iloc[0].to_dict()}")
            logger.info(f"First row keys: {list(df.iloc[0].keys())}")
        
        # Extract transfer information from the DataFrame (first row)
        transfer_info = {
            'vendor': 'Unknown',
            'manifest_id': 'N/A',
            'accepted_date': 'N/A'
        }
        
        if not df.empty:
            first_row = df.iloc[0]
            logger.info(f"Extracting from first row: {first_row.to_dict()}")
            
            # Use the exact column names that exist in the data
            if 'Vendor' in first_row:
                transfer_info['vendor'] = str(first_row['Vendor'])
                logger.info(f"Vendor: {transfer_info['vendor']}")
            
            if 'Barcode*' in first_row:
                transfer_info['manifest_id'] = str(first_row['Barcode*'])
                logger.info(f"Manifest ID (Barcode): {transfer_info['manifest_id']}")
            
            if 'Accepted Date' in first_row:
                transfer_info['accepted_date'] = str(first_row['Accepted Date'])
                logger.info(f"Accepted Date: {transfer_info['accepted_date']}")
        
        logger.info(f"Final transfer info: {transfer_info}")
        
        # Format data for template
        products = []
        for idx, row in df.iterrows():
            product = {
                'id': idx,
                'name': str(row.get('Product Name*', '')),
                'strain': str(row.get('Strain Name', '')),
                'sku': str(row.get('Barcode*', '')),
                'quantity': str(row.get('Quantity Received*', '')),
                'source': format_type or 'Unknown',
                'vendor': str(row.get('Vendor', 'Unknown')),
                'manifest_id': str(row.get('Barcode*', 'N/A')),
                'accepted_date': str(row.get('Accepted Date', 'N/A')),
                'type': str(row.get('Product Type*', 'Unknown')),
                'cost': float(row.get('Cost', 0)) if 'Cost' in row else 0
            }
            products.append(product)

        # Group by weight, sort alphabetically, then split into subgroups of 4
        from collections import defaultdict
        grouped = defaultdict(list)
        for product in products:
            key = str(product.get('weight', 'Unknown'))
            grouped[key].append(product)
        sorted_groups = []
        for group_key in sorted(grouped.keys()):
            group_products = sorted(grouped[group_key], key=lambda x: x['name'].lower())
            # Split group_products into chunks of 4
            for i in range(0, len(group_products), 4):
                sorted_groups.append({
                    'group_label': group_key,
                    'products': group_products[i:i+4]
                })

        # Load configuration
        config = load_config()

        return render_template(
            'data_view.html',
            groups=sorted_groups,
            format_type=format_type,
            theme=config['SETTINGS'].get('theme', 'dark'),
            version=APP_VERSION,
            vendor=transfer_info['vendor'],
            order_date=transfer_info['accepted_date']
        )
    except Exception as e:
        logger.error(f'Error in data_view: {str(e)}', exc_info=True)
        flash('Error loading data. Please try again.')
        return redirect(url_for('index'))

@app.route('/generate-slips', methods=['POST'])
def generate_slips():
    """Generate inventory slips using the original template-based method"""
    try:
        # Get selected products
        selected_indices = request.form.getlist('selected_indices[]')
        
        if not selected_indices:
            flash('No products selected.')
            return redirect(url_for('data_view'))
        
        # Convert indices to integers
        selected_indices = [int(idx) for idx in selected_indices]
        logger.info(f"Selected indices: {selected_indices}")
        
        # Load data from session using chunked data
        df_json = get_chunked_data('df_json')
        
        if df_json is None:
            flash('No data available. Please load data first.')
            return redirect(url_for('index'))
        
        # Convert JSON to DataFrame
        try:
            if isinstance(df_json, list):
                df = pd.DataFrame(df_json)
            else:
                df = pd.read_json(df_json, orient='records')
        except Exception as e:
            logger.error(f"Error converting JSON to DataFrame: {str(e)}")
            flash('Error loading data. Please try again.')
            return redirect(url_for('data_view'))
        
        logger.info(f"DataFrame shape: {df.shape}")
        logger.info(f"DataFrame columns: {df.columns.tolist()}")
        
        # Get only selected rows
        selected_df = df.iloc[selected_indices].copy()
        logger.info(f"Selected DataFrame shape: {selected_df.shape}")
        
        # Load configuration
        config = load_config()
        
        # Generate the file
        status_messages = []
        progress_values = []
        
        def status_callback(msg):
            status_messages.append(msg)
            logger.info(f"Status: {msg}")
        
        def progress_callback(value):
            progress_values.append(value)
        
        logger.info("Starting document generation...")
        success, result = run_full_process_inventory_slips(
            selected_df,
            config,
            status_callback,
            progress_callback
        )
        
        if success:
            logger.info(f"Document generated successfully: {result}")
            # Validate the generated file
            if not validate_docx(result):
                logger.error("Generated document failed validation")
                flash('Generated document appears to be corrupted. Please try again.')
                return redirect(url_for('data_view'))
            
            # Return the file for download
            return send_file(
                result,
                as_attachment=True,
                download_name=os.path.basename(result),
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        else:
            logger.error(f"Document generation failed: {result}")
            flash(f'Failed to generate inventory slips: {result}')
            return redirect(url_for('data_view'))
    
    except Exception as e:
        logger.error(f"Error in generate_slips: {str(e)}", exc_info=True)
        flash(f'Error generating slips: {str(e)}')
        return redirect(url_for('data_view'))

# To this:
@app.route('/generate_robust_slips_docx', methods=['POST'])
def generate_robust_slips_docx():
    """Generate robust inventory slips without complex template rendering"""
    try:
        # Get selected products
        selected_indices = request.form.getlist('selected_indices[]')
        
        if not selected_indices:
            flash('No products selected.')
            return redirect(url_for('data_view'))
        
        # Convert indices to integers
        selected_indices = [int(idx) for idx in selected_indices]
        logger.info(f"Selected indices for robust slip: {selected_indices}")
        
        # Load data from session using chunked data
        df_json = get_chunked_data('df_json')
        
        if df_json is None:
            flash('No data available. Please load data first.')
            return redirect(url_for('data_view'))
        
        # Convert JSON to DataFrame
        try:
            if isinstance(df_json, list):
                df = pd.DataFrame(df_json)
            else:
                df = pd.read_json(df_json, orient='records')
        except Exception as e:
            logger.error(f"Error converting JSON to DataFrame: {str(e)}")
            flash('Error loading data. Please try again.')
            return redirect(url_for('data_view'))
        
        # Get only selected rows
        selected_df = df.iloc[selected_indices].copy()
        logger.info(f"Selected DataFrame shape for robust slip: {selected_df.shape}")
        
        # Load configuration
        config = load_config()
        
        # Generate the robust file
        def status_callback(msg):
            logger.info(f"Robust slip status: {msg}")
        
        logger.info("Starting robust document generation...")
        success, result = create_robust_inventory_slip(
            selected_df,
            config,
            status_callback
        )
        
        if success:
            logger.info(f"Robust document generated successfully: {result}")
            # Return the file for download
            return send_file(
                result,
                as_attachment=True,
                download_name=os.path.basename(result),
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        else:
            logger.error(f"Robust document generation failed: {result}")
            flash(f'Failed to generate robust inventory slips: {result}')
            return redirect(url_for('data_view'))
    
    except Exception as e:
        logger.error(f"Error in generate_robust_slips: {str(e)}", exc_info=True)
        flash(f'Error generating robust slips: {str(e)}')
        return redirect(url_for('data_view'))

@app.route('/show-result')
def show_result():
    # Get output file path from session
    output_file = session.get('output_file', None)
    
    if not output_file or not os.path.exists(output_file):
        flash('No output file available.')
        return redirect(url_for('index'))
    
    # Get filename for display
    filename = os.path.basename(output_file)
    
    # Load configuration
    config = load_config()
    
    return render_template(
        'result.html',
        filename=filename,
        theme=config['SETTINGS'].get('theme', 'dark'),
        version=APP_VERSION
    )

@app.route('/download-file')
def download_file():
    # Get output file path from session
    output_file = session.get('output_file', None)
    
    if not output_file or not os.path.exists(output_file):
        flash('No output file available.')
        return redirect(url_for('index'))
    
    # Return the file for download
    return send_file(output_file, as_attachment=True)

@app.route('/settings', methods=['GET', 'POST'])
def settings():
    config = load_config()
    
    if request.method == 'POST':
        # Update settings from form
        if 'items_per_page' in request.form:
            config['SETTINGS']['items_per_page'] = request.form['items_per_page']
        
        if 'theme' in request.form:
            config['SETTINGS']['theme'] = request.form['theme']
        
        if 'api_key' in request.form:
            if 'API' not in config:
                config['API'] = {}
            config['API']['bamboo_key'] = request.form['api_key']
        
        if 'outputDir' in request.form:
            output_dir = request.form['outputDir']
            if output_dir:
                config['PATHS']['output_dir'] = output_dir
        
        # Save updated config
        save_config(config)
        flash('Settings saved successfully')
        return redirect(url_for('index'))
    
    return render_template(
        'settings.html',
        config=config,
        theme=config['SETTINGS'].get('theme', 'dark'),
        version=APP_VERSION
    )

def require_api_key(f):
    """Decorator to require API key for certain routes"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        config = load_config()
        api_type = request.args.get('api_type', 'bamboo')
        
        if 'API' not in config or f'{api_type}_key' not in config['API']:
            flash(f'No API key configured for {api_type}. Please add it in settings.')
            return redirect(url_for('settings'))
        return f(*args, **kwargs)
    return decorated_function

class APIClient:
    def __init__(self, api_type, config):
        self.api_type = api_type
        self.config = config
        self.api_config = API_CONFIGS.get(api_type)
        if not self.api_config:
            raise ValueError(f"Unsupported API type: {api_type}")
        
    def get_headers(self):
        """Get headers for API request based on API type"""
        headers = {
            'User-Agent': f'InventorySlipGenerator/{APP_VERSION}',
            'Accept': 'application/json',
            'Content-Type': 'application/json'
        }
        
        api_key = self.config['API'].get(f'{self.api_type}_key')
        if not api_key:
            return headers
            
        if self.api_config['auth_type'] == 'bearer':
            headers['Authorization'] = f'Bearer {api_key}'
        elif self.api_config['auth_type'] == 'basic':
            encoded = base64.b64encode(f'{api_key}:'.encode()).decode()
            headers['Authorization'] = f'Basic {encoded}'
        
        return headers
    
    def make_request(self, endpoint, method='GET', params=None, data=None):
        """Make API request with proper error handling"""
        url = f"{self.api_config['base_url']}/{self.api_config['version']}/{endpoint}"
        headers = self.get_headers()
        
        try:
            response = requests.request(
                method=method,
                url=url,
                headers=headers,
                params=params,
                json=data,
                timeout=30
            )
            response.raise_for_status()
            return response.json()
        except requests.exceptions.RequestException as e:
            logger.error(f"API request failed: {str(e)}")
            raise

# Add these new routes

@app.route('/api/fetch-transfers', methods=['POST'])
@require_api_key
def fetch_transfers():
    """Fetch transfer data from selected API"""
    try:
        api_type = request.form.get('api_type', 'bamboo')
        date_from = request.form.get('date_from')
        date_to = request.form.get('date_to')
        
        config = load_config()
        client = APIClient(api_type, config)
        
        # Fetch data based on API type
        if api_type == 'bamboo':
            data = client.make_request('transfers', params={'start_date': date_from, 'end_date': date_to})
            result_df = parse_bamboo_data(data)
        elif api_type == 'cultivera':
            data = client.make_request('manifests', params={'fromDate': date_from, 'toDate': date_to})
            result_df = parse_cultivera_data(data)
        elif api_type == 'growflow':
            data = client.make_request('inventory/transfers', params={'dateStart': date_from, 'dateEnd': date_to})
            result_df = parse_growflow_data(data)
        else:
            return jsonify({'error': 'Unsupported API type'}), 400
        
        if result_df is None or result_df.empty:
            return jsonify({'error': 'No data found'}), 404
            
        # Store in session
        session['df_json'] = result_df.to_json(orient='records')
        session['format_type'] = api_type
        session['raw_json'] = json.dumps(data)
        
        return jsonify({
            'success': True,
            'message': f'Successfully fetched {len(result_df)} records',
            'redirect': url_for('data_view')
        })
        
    except Exception as e:
        logger.error(f"API fetch error: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/validate-key', methods=['POST'])
def validate_api_key():
    """Validate API key for selected service"""
    try:
        api_type = request.form.get('api_type')
        api_key = request.form.get('api_key')
        
        if not api_type or not api_key:
            return jsonify({'valid': False, 'message': 'Missing required parameters'}), 400
            
        config = load_config()
        client = APIClient(api_type, {'API': {f'{api_type}_key': api_key}})
        
        if 'API' not in config:
            config['API'] = {}
        config['API'][f'{api_type}_key'] = api_key
        save_config(config)
        
        return jsonify({
            'valid': True,
            'message': f'{api_type.title()} API key validated and saved'
        })
        
    except requests.exceptions.HTTPError as e:
        if e.response.status_code == 401:
            return jsonify({
                'valid': False,
                'message': 'Invalid API key'
            }), 401
        return jsonify({
            'valid': False,
            'message': f'API error: {str(e)}'
        }), e.response.status_code
    except Exception as e:
        return jsonify({
            'valid': False,
            'message': f'Validation error: {str(e)}'
        }), 500
        

@app.route('/api/settings', methods=['GET', 'POST'])
def api_settings():
    """Manage API settings"""
    config = load_config()
    
    if request.method == 'POST':
        api_type = request.form.get('api_type')
        api_key = request.form.get('api_key')
        
        if api_type and api_key:
            if 'API' not in config:
                config['API'] = {}
            config['API'][f'{api_type}_key'] = api_key
            save_config(config)
            
            flash(f'{api_type.title()} API key updated successfully')
            return redirect(url_for('settings'))
            
    # Get current API keys
    api_keys = {
        'bamboo': config.get('API', {}).get('bamboo_key', ''),
        'cultivera': config.get('API', {}).get('cultivera_key', ''),
        'growflow': config.get('API', {}).get('growflow_key', '')
    }
    
    return render_template(
        'api_settings.html',
        api_keys=api_keys,
                                                         theme=config['SETTINGS'].get('theme', 'dark'),
        version=APP_VERSION
    )

# Add these error handlers

class APIError(Exception):
    """Base class for API-related errors"""
    pass

class APIAuthError(APIError):
    """Authentication error"""
    pass

class APIRateLimit(APIError):
    """Rate limit exceeded"""
    pass

class APIDataError(APIError):
    """Data processing error"""
    pass

@app.errorhandler(APIError)
def handle_api_error(error):
    """Handle API errors gracefully"""
    if isinstance(error, APIAuthError):
        flash(f'Authentication error: {error}', 'error')
    elif isinstance(error, APIRateLimit):
        flash(f'Rate limit exceeded: {error}', 'warning')
    elif isinstance(error, APIDataError):
        flash(f'Data error: {error}', 'error')
    else:
        flash(f'API error: {error}', 'error')
    return redirect(url_for('api_settings'))

@app.route('/test-chunked-data')
def test_chunked_data():
    """Test route to debug chunked data storage"""
    try:
        # Test storing and retrieving data
        test_data = {"test": "data", "number": 123}
        logger.info("Testing chunked data storage...")
        
        store_chunked_data('test_key', json.dumps(test_data))
        retrieved_data = get_chunked_data('test_key')
        
        logger.info(f"Test data: {test_data}")
        logger.info(f"Retrieved data: {retrieved_data}")
        
        # Check session contents
        session_keys = [k for k in session.keys() if k.startswith('test_key')]
        logger.info(f"Session keys for test_key: {session_keys}")
        
        return jsonify({
            'success': True,
            'original': test_data,
            'retrieved': retrieved_data,
            'session_keys': session_keys
        })
    except Exception as e:
        logger.error(f"Test failed: {str(e)}", exc_info=True)
        return jsonify({'success': False, 'error': str(e)})

@app.route('/debug-session')

def debug_session():
    """Debug session issues for Chrome compatibility"""
    try:
        # Get session info
        session_info = {
            'session_id': session.get('_id', 'Not set'),
            'session_keys': list(session.keys()),
            'session_size': len(str(session)),
            'user_agent': request.headers.get('User-Agent', 'Unknown'),
            'chrome_version': None
        }
        
        # Try to detect Chrome version
        user_agent = request.headers.get('User-Agent', '')
        if 'Chrome/' in user_agent:
            try:
                chrome_version = user_agent.split('Chrome/')[1].split(' ')[0]
                session_info['chrome_version'] = chrome_version
            except:
                session_info['chrome_version'] = 'Unknown'
        
        # Test session storage
        test_key = 'chrome_test'
        session[test_key] = 'test_value'
        session_info['test_storage'] = session.get(test_key) == 'test_value'
        
        return jsonify({
            'success': True,
            'session_info': session_info,
            'headers': dict(request.headers)
        })
    except Exception as e:
        logger.error(f"Session debug failed: {str(e)}", exc_info=True)
        return jsonify({'success': False, 'error': str(e)})

@app.route('/about')
def about():
    """About page"""
    config = load_config()
    return render_template('about.html', config=config, version=APP_VERSION)

@app.route('/')
def index():
    # Load configuration for the template
    config = load_config()
    return render_template('index.html', config=config, version=APP_VERSION)

@app.route('/', methods=['OPTIONS'])
def handle_options():
    """Handle OPTIONS requests for CORS preflight"""
    response = app.make_default_options_response()
    response.headers['Access-Control-Allow-Origin'] = 'http://localhost:*'
    response.headers['Access-Control-Allow-Methods'] = 'GET, POST, PUT, DELETE, OPTIONS'
    response.headers['Access-Control-Allow-Headers'] = 'Content-Type, Authorization'
    response.headers['Access-Control-Allow-Credentials'] = 'true'
    return response

@app.route('/search_json_or_api', methods=['POST'])
def search_json_or_api():
    """Handle both JSON paste and API URL functionality"""
    try:
        search_input = request.form.get('search_input', '').strip()
        
        if not search_input:
            flash('Please enter JSON data or an API URL.')
        search_input = request.form.get('search_input', '').strip()
        
        if not search_input:
            flash('Please enter JSON data or an API URL.')
            return redirect(url_for('index'))
        
        # Check if it looks like a URL
        if search_input.startswith(('http://', 'https://')):
            # Handle as API URL
            return load_from_url(search_input)
        else:
            # Handle as JSON data
            return paste_json_data(search_input)
            
    except Exception as e:
        logger.error(f"Error in search_json_or_api: {str(e)}")
        flash(f'Error processing input: {str(e)}')
        return redirect(url_for('index'))

def paste_json_data(json_text):
    """Handle JSON data processing"""
    try:
        if not json_text.strip():
            flash('Please enter JSON data.')
            return redirect(url_for('index'))
        
        # Parse JSON
        try:
            json_data = json.loads(json_text)
        except json.JSONDecodeError as e:
            flash(f'Invalid JSON format: {str(e)}')
            return redirect(url_for('index'))
        
        # Determine format and process
        if 'transfers' in json_data:
            # Bamboo format
            df = parse_bamboo_data(json_data)
            format_type = 'Bamboo'
        elif 'manifest' in json_data:
            # Cultivera format
            df = parse_cultivera_data(json_data)
            format_type = 'Cultivera'
        elif 'data' in json_data and isinstance(json_data['data'], list):
            # GrowFlow format
            df = parse_growflow_data(json_data)
            format_type = 'GrowFlow'
        else:
            # Generic inventory format
            df = parse_inventory_json(json_data)
            format_type = 'Generic'
        
        if df is None or df.empty:
            flash('No valid data found in JSON.')
            return redirect(url_for('index'))
        
        # Store data in session
        store_chunked_data('df_json', df)
        session['format_type'] = format_type
        
        flash(f'Successfully loaded {len(df)} items from {format_type} data.')
        return redirect(url_for('data_view'))
        
    except Exception as e:
        logger.error(f"Error in paste_json_data: {str(e)}")
        flash(f'Error processing JSON data: {str(e)}')
        return redirect(url_for('index'))

@app.route('/open_downloads', methods=['GET'])
def open_downloads():
    """Open the downloads folder"""
    try:
        config = load_config()
        downloads_path = config['PATHS']['output_dir']
        
        # Open the folder
        open_file(downloads_path)
        
        return jsonify({'success': True, 'message': 'Downloads folder opened'})
    except Exception as e:
        logger.error(f"Error opening downloads folder: {str(e)}")
        return jsonify({'success': False, 'message': str(e)})

def validate_docx(file_path):
    """Validate that a Word document is readable and not corrupted"""
    try:
        if not os.path.exists(file_path):
            return False
        
        file_size = os.path.getsize(file_path)
        if file_size < 1000:  # Too small to be a valid Word document
            return False
        
        # Try to open and read the document
        doc = Document(file_path)
        
        # Check if document has any content
        if len(doc.paragraphs) == 0 and len(doc.tables) == 0:
            return False
        
        return True
        
    except Exception as e:
        logger.error(f"Document validation failed for {file_path}: {str(e)}")
        return False

@app.route('/clear_data')
def clear_data():
    """Clear all session data"""
    try:
        # Clear all chunked data
        for key in ['df_json', 'raw_json']:
            clear_chunked_data(key)
        
        # Clear other session data
        session.pop('format_type', None)
        
        flash('All data has been cleared successfully.', 'success')
        return redirect(url_for('index'))
    except Exception as e:
        logger.error(f"Error clearing data: {str(e)}")
        flash('Error clearing data. Please try again.', 'error')
        return redirect(url_for('index'))

@app.route('/select_directory', methods=['POST'])
def select_directory():
    """Select output directory"""
    try:
        # For now, return the default downloads directory
        # In a full implementation, this would open a file dialog
        config = load_config()
        downloads_path = config['PATHS']['output_dir']
        
        return jsonify({
            'success': True,
            'directory': downloads_path,
            'message': f'Using default directory: {downloads_path}'
        })
    except Exception as e:
        logger.error(f"Error selecting directory: {str(e)}")
        return jsonify({'success': False, 'message': str(e)})

@app.route('/fetch_api', methods=['POST'])
def fetch_api():
    """Fetch data from API"""
    try:
        # This would handle API data fetching
        # For now, redirect to the search_json_or_api route
        return redirect(url_for('search_json_or_api'))
    except Exception as e:
        logger.error(f"Error fetching API data: {str(e)}")
        flash(f'Error fetching API data: {str(e)}', 'error')
        return redirect(url_for('index'))

def ensure_ocr_column():
    conn = sqlite3.connect(PDF_DB_PATH)
    c = conn.cursor()
    # Check if ocr_text column exists
    c.execute("PRAGMA table_info(pdf_inventory)")
    columns = [row[1] for row in c.fetchall()]
    if 'ocr_text' not in columns:
        c.execute('ALTER TABLE pdf_inventory ADD COLUMN ocr_text TEXT')
        conn.commit()
    conn.close()

# Call this before any insert/update to pdf_inventory
ensure_ocr_column()

@app.route('/download_pdf/<filename>')
def download_pdf(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename, as_attachment=True)

# Route to delete a PDF scan
@app.route('/delete_pdf/<filename>', methods=['POST'])
def delete_pdf(filename):
    # Remove from DB
    conn = sqlite3.connect(PDF_DB_PATH)
    c = conn.cursor()
    c.execute('DELETE FROM pdf_inventory WHERE filename = ?', (filename,))
    conn.commit()
    conn.close()
    # Remove file from uploads folder
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    if os.path.exists(file_path):
        os.remove(file_path)
    flash(f'Deleted {filename}')
    return redirect(url_for('list_pdfs'))

@app.route('/download_excel')
def download_excel():
    # Fetch PDF metadata and OCR text from DB
    conn = sqlite3.connect(PDF_DB_PATH)
    c = conn.cursor()
    c.execute('SELECT filename, upload_date, ocr_text FROM pdf_inventory ORDER BY upload_date DESC')
    pdfs = c.fetchall()
    conn.close()
    # Parse OCR text into structured columns
    def parse_slip(text):
        import re
        lines = text.splitlines()
        filtered = []
        for line in lines:
            if re.search(r'(\d{4}-\d{2}-\d{2}|JSM LLC|Only B\'s|Dank Czar|Flavour Bar|Omega Distillate|Medically Compliant|SKU:|Initial Qty Issued:|Qty Received:)', line):
                filtered.append(line.strip())
        filtered_text = '\n'.join(filtered)
        date = re.search(r'\d{4}-\d{2}-\d{2}', filtered_text)
        vendor = re.search(r'JSM LLC|Only B\'s|Dank Czar|Flavour Bar|Omega Distillate', filtered_text)
        product = re.search(r'(Medically Compliant.*?)(SKU:|$)', filtered_text, re.DOTALL)
        sku = re.search(r'SKU:\s*(\d+)', filtered_text)
        qty_issued = re.search(r'Initial Qty Issued:\s*\|?\s*(\d+)?', filtered_text)
        qty_received = re.search(r'Qty Received:\s*\|?\s*(\d+)?', filtered_text)
        return {
            'date': date.group(0) if date else '',
            'vendor': vendor.group(0) if vendor else '',
            'product': product.group(1).replace('\n', ' ').strip() if product else '',
            'sku': sku.group(1) if sku else '',
            'qty_issued': qty_issued.group(1) if qty_issued else '',
            'qty_received': qty_received.group(1) if qty_received else ''
        }
    rows = []
    for filename, upload_date, ocr_text in pdfs:
        slip = parse_slip(ocr_text or '')
        rows.append({
            'Filename': filename,
            'Upload Date': upload_date,
            'Date': slip['date'],
            'Vendor': slip['vendor'],
            'Product': slip['product'],
            'SKU': slip['sku'],
            'Initial Qty Issued': slip['qty_issued'],
            'Qty Received': slip['qty_received']
        })
    import pandas as pd
    df = pd.DataFrame(rows)
    from io import BytesIO
    output = BytesIO()
    df.to_excel(output, index=False)
    output.seek(0)
    # Use product name from first slip for filename
    excel_name = 'scanned_pdfs.xlsx'
    if len(rows) > 0 and rows[0]['Vendor']:
        safe_vendor_name = re.sub(r'[^A-Za-z0-9_\-]', '_', rows[0]['Vendor'])[:40]
        excel_name = f"{safe_vendor_name}_inventory.xlsx"
    return send_file(output, download_name=excel_name, as_attachment=True, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

if __name__ == '__main__':
    try:
        # Clean up any temporary files from previous runs
        cleanup_temp_files()
        
        # Try different ports in case default is taken
        ports = [8000, 8001, 8080, 8081, 8888, 9000]
        
        for port in ports:
            try:
                print(f"Attempting to start server on port {port}...")
                
                # Open browser with more reliable method
                def open_browser():
                    try:
                        # Try Chrome first with --new-window flag
                        chrome_path = ''
                        if sys.platform == "darwin":  # macOS
                            chrome_path = 'open -a /Applications/Google\ Chrome.app %s'
                        elif sys.platform == "win32":  # Windows
                            chrome_path = 'C:/Program Files/Google/Chrome/Application/chrome.exe %s'
                        
                        url = f'http://localhost:{port}'
                        
                        # Add Chrome flags to handle authentication issues
                        chrome_flags = [
                            '--disable-web-security',
                            '--disable-features=VizDisplayCompositor',
                            '--disable-site-isolation-trials',
                            '--disable-features=TranslateUI',
                            '--disable-ipc-flooding-protection',
                            '--no-first-run',
                            '--no-default-browser-check',
                            '--disable-default-apps'
                        ]
                        
                        if chrome_path:
                            # Try to open with Chrome flags
                            try:
                                import subprocess
                                if sys.platform == "darwin":  # macOS
                                    cmd = ['open', '-a', 'Google Chrome', '--args'] + chrome_flags + [url]
                                elif sys.platform == "win32":  # Windows
                                    cmd = ['chrome.exe'] + chrome_flags + [url]
                                else:
                                    cmd = ['google-chrome'] + chrome_flags + [url]
                                
                                subprocess.Popen(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                                print(f"Opened Chrome with authentication-friendly flags")
                            except Exception as e:
                                print(f"Failed to open Chrome with flags: {e}")
                                # Fallback to webbrowser
                                webbrowser.get(chrome_path).open(url, new=2)
                        else:
                            webbrowser.open(url, new=2)
                    except Exception as e:
                        print(f"Error opening browser: {e}")
                        # Fallback to default browser
                        webbrowser.open(f'http://localhost:{port}', new=2)

                # Delay browser opening slightly
                threading.Timer(2.0, open_browser).start()
                
                app.run(
                    host='localhost',
                    port=port,
                    debug=True,
                    use_reloader=False  # Prevent duplicate browser windows
                )
                break  # If server starts successfully, break the loop
                
            except OSError as e:
                print(f"Port {port} is in use, trying next port...")
                if port == ports[-1]:  # If we've tried all ports
                    print("Could not find an available port. Please try again or manually specify a port.")
                    print("Available ports to try manually: 5001, 5002, 5003, 8000, 8001, 8080, 8081, 8888, 9000")
                continue
                
    except Exception as e:
        print(f"Failed to start server: {str(e)}")
        sys.exit(1)
