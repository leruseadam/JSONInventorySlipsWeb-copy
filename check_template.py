from docx import Document
import os
from xml.etree import ElementTree as ET
import zipfile

def extract_text_from_xml(xml_string):
    """Extract text content from Word XML"""
    try:
        root = ET.fromstring(xml_string)
        # Find all w:t elements (text runs in Word)
        text_elements = root.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
        return [elem.text for elem in text_elements if elem.text]
    except ET.ParseError:
        return []

def check_docx(filepath):
    print(f"\nChecking {os.path.basename(filepath)}...")
    try:
        # Try reading with python-docx first
        doc = Document(filepath)
        print("\nParagraphs:")
        for p in doc.paragraphs:
            if p.text.strip():
                print(f"- {p.text}")
                for run in p.runs:
                    print(f"  Run: '{run.text}'")
        
        print("\nTables:")
        for table in doc.tables:
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    if any(p.text.strip() for p in cell.paragraphs):
                        print(f"- Cell [{i},{j}]:")
                        for p in cell.paragraphs:
                            if p.text.strip():
                                print(f"  {p.text}")
                                for run in p.runs:
                                    print(f"    Run: '{run.text}'")
                                    
        print("\nRaw XML Content:")
        # Also try reading the raw XML
        with zipfile.ZipFile(filepath) as docx:
            # Read main document content
            with docx.open('word/document.xml') as xml_content:
                xml_str = xml_content.read().decode('utf-8')
                text_elements = extract_text_from_xml(xml_str)
                if text_elements:
                    print("\nFound text in document.xml:")
                    for text in text_elements:
                        print(f"- '{text}'")
                else:
                    print("No text found in document.xml")
                    
            # Also check headers and footers
            for item in docx.namelist():
                if item.startswith('word/header') or item.startswith('word/footer'):
                    with docx.open(item) as xml_content:
                        xml_str = xml_content.read().decode('utf-8')
                        text_elements = extract_text_from_xml(xml_str)
                        if text_elements:
                            print(f"\nFound text in {item}:")
                            for text in text_elements:
                                print(f"- '{text}'")
                                
    except Exception as e:
        print(f"Error: {e}")

for filename in [
    "templates/documents/InventorySlips.docx",
    "templates/documents/InventorySlips.backup.docx",
    "templates/documents/InventorySlips_old.docx",
    "templates/documents/template_check.docx"
]:
    check_docx(filename)
